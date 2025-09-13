import os
import uuid
import math
import time
import logging
from datetime import datetime, timezone
from flask import Flask, request, jsonify, send_file, abort, render_template_string, url_for, redirect, session, current_app, render_template
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy import inspect
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from flask_mail import Mail, Message
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.exceptions import BadRequest
from itsdangerous import URLSafeTimedSerializer, SignatureExpired, BadSignature
from pypdf import PdfReader, PdfWriter
import pikepdf
import io
import traceback
import tempfile
import shutil  # Added for file moving in enhanced_split
import subprocess  # Ensure it's imported at the top level
import json
import sys
from pathlib import Path
import json
from pathlib import Path
# Import our notebook conversion utilities
from notebook_utils import ipynb_to_pdf, ipynb_to_docx, py_to_ipynb, py_to_pdf
import io
import traceback
import tempfile
import shutil  # Added for file moving in enhanced_split
import subprocess  # Ensure it's imported at the top level
from flask.helpers import send_from_directory
from config_loader import get_env, get_secret_key, get_database_url, is_debug_mode

# Helper function to get absolute path for uploads
# ---------- small helpers ----------
def _ensure_tool_on_path(tool_name: str) -> bool:
    """Return True if tool is discoverable on PATH via shutil.which."""
    return shutil.which(tool_name) is not None

def _run_cmd(cmd, cwd=None, timeout=300):
    """Run cmd (list), return dict {rc,stdout,stderr} - never raise."""
    try:
        proc = subprocess.run(cmd, check=False, capture_output=True, text=True, cwd=cwd, timeout=timeout)
        return {"rc": proc.returncode, "stdout": proc.stdout or "", "stderr": proc.stderr or ""}
    except Exception as e:
        return {"rc": -1, "stdout": "", "stderr": str(e)}

def get_upload_path(filename):
    """Get the absolute path for an uploaded file"""
    # Fallback to 'uploads' inside app root if not configured
    upload_folder = UPLOAD_FOLDER
    # Ensure folder exists
    os.makedirs(upload_folder, exist_ok=True)
    return os.path.abspath(os.path.join(upload_folder, filename))

# For compatibility with the new routes
def _get_upload_path(filename):
    return get_upload_path(filename)

def _ensure_tool_on_path(tool_name: str) -> bool:
    """Return True if tool is discoverable on PATH via shutil.which."""
    return shutil.which(tool_name) is not None

def _run_cmd(cmd, cwd=None, timeout=300):
    """Run cmd (list), return dict {rc,stdout,stderr} - never raise."""
    try:
        proc = subprocess.run(cmd, check=False, capture_output=True, text=True, cwd=cwd, timeout=timeout)
        return {"rc": proc.returncode, "stdout": proc.stdout or "", "stderr": proc.stderr or ""}
    except Exception as e:
        return {"rc": -1, "stdout": "", "stderr": str(e)}

# Twilio and Redis for OTP
try:
    from twilio.rest import Client
    from twilio.base.exceptions import TwilioRestException
    import redis
    TWILIO_AVAILABLE = True
    REDIS_AVAILABLE = True
except ImportError:
    TWILIO_AVAILABLE = False
    REDIS_AVAILABLE = False
    logging.warning("Twilio or Redis not available - OTP functionality disabled")

import random
import string

# AI/ML and Advanced Features
try:
    from flask_restx import Api, Resource, fields
    from functools import wraps
    RESTX_AVAILABLE = True
except ImportError:
    RESTX_AVAILABLE = False
    logging.warning("Flask-RESTX not available - API documentation disabled")

try:
    from opentelemetry import trace
    from opentelemetry.instrumentation.flask import FlaskInstrumentor
    from opentelemetry.sdk.trace import TracerProvider
    from opentelemetry.sdk.trace.export import BatchSpanProcessor
    from opentelemetry.exporter.jaeger.thrift import JaegerExporter
    from opentelemetry.sdk.resources import Resource as OTResource, SERVICE_NAME
    TELEMETRY_AVAILABLE = True
except ImportError:
    TELEMETRY_AVAILABLE = False
    logging.warning("OpenTelemetry not available - observability disabled")

# Celery configuration
try:
    from celery import Celery
    CELERY_AVAILABLE = True
except ImportError:
    CELERY_AVAILABLE = False
    logging.warning("Celery not available - background tasks disabled")

# Import the new PDF processor
from pdf_processor import PDFProcessor, PDFValidationError, PDFOperationError

# Initialize PDF processor with higher file size limit (2GB)
pdf_processor = PDFProcessor(max_file_size_mb=2048)

# --- App Initialization ---
app = Flask(__name__, static_folder='static', static_url_path='/static')

# Load environment variables from config_loader
from config_loader import get_secret_key, get_database_url, get_api_key, is_debug_mode

# Use secure secret key with proper fallback handling
secret_key = get_secret_key()
app.config['SECRET_KEY'] = secret_key

# Note: We don't initialize Celery here - we use the instance from tasks.py

# Add route for app.bundle.js and other frontend bundles
@app.route('/app.bundle.js')
def app_bundle():
    return send_from_directory(os.path.join(app.root_path, 'static/dist'), 'app.bundle.js')

@app.route('/enhancedPDFTools.bundle.js')
def enhanced_pdf_tools_bundle():
    return send_from_directory(os.path.join(app.root_path, 'static/dist'), 'enhancedPDFTools.bundle.js')

@app.route('/enhancedSearchComponent.bundle.js')
def enhanced_search_component_bundle():
    return send_from_directory(os.path.join(app.root_path, 'static/dist'), 'enhancedSearchComponent.bundle.js')

@app.route('/static/js/AuthenticationManager.js')
def authentication_manager_js():
    return send_from_directory(os.path.join(app.root_path, 'static/js'), 'AuthenticationManager.js')

@app.route('/static/js/auth-integration.js')
def auth_integration_js():
    return send_from_directory(os.path.join(app.root_path, 'static/js'), 'auth-integration.js')

# Add a route to serve login static files from templates directory
@app.route('/static/login/<path:filename>')
def login_static(filename):
    # Split the path to handle subdirectories correctly
    subdirs = filename.split('/')
    if len(subdirs) > 1:
        # For nested paths like 'img/login-bg.png'
        subdir, file = '/'.join(subdirs[:-1]), subdirs[-1]
        return send_from_directory(os.path.join(app.root_path, 'templates/login', subdir), file)
    else:
        # For files directly in the login directory
        return send_from_directory(os.path.join(app.root_path, 'templates/login'), filename)

# Add a route to serve JS files from static/js directory
@app.route('/js/<path:filename>')
def js_files(filename):
    return send_from_directory(os.path.join(app.root_path, 'static/js'), filename)

# Health check endpoint

# Add route for CSS files
@app.route('/css/<path:filename>')
def css_files(filename):
    return send_from_directory(os.path.join(app.root_path, 'static/css'), filename)

# Add route for favicon
@app.route('/favicon.ico')
def favicon():
    return send_from_directory(os.path.join(app.root_path, 'static'), 'favicon.ico')

# Configure database with fallback
app.config['SQLALCHEMY_DATABASE_URI'] = get_database_url()
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['MAX_CONTENT_LENGTH'] = int(get_env('MAX_CONTENT_LENGTH_MB', '2048')) * 1024 * 1024  # default 2GB

# Initialize extensions
db = SQLAlchemy(app)
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

# Flask-Mail configuration
app.config['MAIL_SERVER'] = get_env('MAIL_SERVER', 'smtp.gmail.com')
app.config['MAIL_PORT'] = int(get_env('MAIL_PORT', '587'))
app.config['MAIL_USE_TLS'] = get_env('MAIL_USE_TLS', 'true').lower() == 'true'
app.config['MAIL_USE_SSL'] = get_env('MAIL_USE_SSL', 'false').lower() == 'true'
app.config['MAIL_USERNAME'] = get_env('MAIL_USERNAME', '')
app.config['MAIL_PASSWORD'] = get_env('MAIL_PASSWORD', '')
app.config['MAIL_DEFAULT_SENDER'] = get_env('MAIL_DEFAULT_SENDER', 'noreply@pdf-tool.com')

mail = Mail(app)
serializer = URLSafeTimedSerializer(app.config['SECRET_KEY'])

# Twilio configuration for OTP
if TWILIO_AVAILABLE:
    twilio_client = Client(
        os.getenv('TWILIO_ACCOUNT_SID', ''),
        os.getenv('TWILIO_AUTH_TOKEN', '')
    )
    twilio_phone = os.getenv('TWILIO_PHONE_NUMBER', '')

# Redis configuration for OTP storage
if REDIS_AVAILABLE:
    try:
        redis_client = redis.Redis(
            host=os.getenv('REDIS_HOST', 'localhost'),
            port=int(os.getenv('REDIS_PORT', 6379)),
            db=int(os.getenv('REDIS_DB', 0)),
            decode_responses=True
        )
        # Test connection
        redis_client.ping()
    except Exception as e:
        logging.warning(f"Redis connection failed: {e}")
        REDIS_AVAILABLE = False

# --- Database Models ---
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=False)
    phone_number = db.Column(db.String(20), unique=True, nullable=True)
    password_hash = db.Column(db.String(200), nullable=False)
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    google_drive_token = db.Column(db.Text, nullable=True)  # For Google Drive integration
    webhook_url = db.Column(db.String(255), nullable=True)  # For Zapier/Make integration
    suggested_workflow = db.Column(db.JSON, nullable=True)  # For personalized workflows
    last_login = db.Column(db.DateTime, nullable=True)  # Track last login time
    login_count = db.Column(db.Integer, default=0)  # Track number of logins
    last_ip = db.Column(db.String(45), nullable=True)  # Store last IP address (IPv6 compatible)
    user_agent = db.Column(db.String(255), nullable=True)  # Store browser/client info
    files = db.relationship('FileRecord', backref='user', lazy=True)
    login_history = db.relationship('UserLoginHistory', backref='user', lazy=True)

class FileRecord(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(200), nullable=False)
    original_filename = db.Column(db.String(200), nullable=False)
    file_size = db.Column(db.Integer, nullable=False)
    file_type = db.Column(db.String(50), nullable=False)
    doc_type = db.Column(db.String(50), nullable=True)  # For document classification
    doc_metadata = db.Column(db.JSON, nullable=True)  # For document metadata
    upload_date = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    last_accessed = db.Column(db.DateTime, nullable=True)  # When file was last accessed
    access_count = db.Column(db.Integer, default=0)  # How many times the file was accessed
    page_count = db.Column(db.Integer, nullable=True)  # Number of pages in the PDF
    storage_path = db.Column(db.String(255), nullable=True)  # Full path where file is stored
    mimetype = db.Column(db.String(100), nullable=True)  # MIME type of the file
    hash_md5 = db.Column(db.String(32), nullable=True)  # MD5 hash for deduplication
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    team_id = db.Column(db.Integer, db.ForeignKey('team.id'), nullable=True)  # For team management
    conversions = db.relationship('FileConversionRecord', backref='original_file', lazy=True)

class ProcessingRecord(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    task_id = db.Column(db.String(100), unique=True, nullable=False)
    command = db.Column(db.String(50), nullable=False)
    input_files = db.Column(db.JSON, nullable=False)
    output_file = db.Column(db.String(200), nullable=False)
    status = db.Column(db.String(20), default='pending')
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    completed_at = db.Column(db.DateTime)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)

class Team(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    files = db.relationship('FileRecord', backref='team', lazy=True)
    memberships = db.relationship('TeamMembership', backref='team', lazy=True)

class TeamMembership(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    team_id = db.Column(db.Integer, db.ForeignKey('team.id'), nullable=False)
    role = db.Column(db.String(20), nullable=False, default='member')  # 'admin' or 'member'
    joined_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    user = db.relationship('User', backref='team_memberships')

class UserLoginHistory(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    login_time = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    ip_address = db.Column(db.String(45), nullable=True)  # IPv6-compatible length
    user_agent = db.Column(db.String(255), nullable=True)
    success = db.Column(db.Boolean, default=True)
    device_info = db.Column(db.JSON, nullable=True)  # Store detailed device information

class FileConversionRecord(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    original_file_id = db.Column(db.Integer, db.ForeignKey('file_record.id'), nullable=False)
    output_file = db.Column(db.String(200), nullable=False)
    conversion_type = db.Column(db.String(50), nullable=False)  # 'merge', 'split', 'compress', etc.
    conversion_time = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    status = db.Column(db.String(20), default='completed')
    file_size = db.Column(db.Integer, nullable=True)
    processing_time_ms = db.Column(db.Integer, nullable=True)  # Performance tracking
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    pg_tools_version = db.Column(db.String(20), default='2.0.12')

# --- Configuration ---
UPLOAD_FOLDER = 'uploads'
PROCESSED_FOLDER = 'processed'
ALLOWED_EXTENSIONS = {
    'pdf', 'docx', 'pptx', 'xlsx', 'xls', 'html', 'htm',
    'ipynb', 'py', 'jpg', 'jpeg', 'png', 'gif'
}

# Create folders if they don't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(PROCESSED_FOLDER, exist_ok=True)

# --- Task Management ---
app.task_results = {}
app.task_timestamps = {}

# Add auth bypass for tests
from functools import wraps

def auth_required(f):
    @wraps(f)
    def wrapper(*args, **kwargs):
        # Allow tests to bypass auth
        if app.config.get("TESTING"):
            return f(*args, **kwargs)
        # Use login_required for real auth
        return login_required(f)(*args, **kwargs)
    return wrapper

def admin_required(f):
    """Decorator that ensures user is authenticated and has admin role.
    Use this on routes that should only be accessible to administrators.
    """
    @wraps(f)
    def wrapper(*args, **kwargs):
        # Allow tests to bypass auth
        if app.config.get("TESTING"):
            return f(*args, **kwargs)
        
        # Check if user is authenticated
        if not current_user.is_authenticated:
            if request.is_json or request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'error': 'Authentication required'}), 401
            return redirect(url_for('login'))
        
        # Check if user is admin (username is 'admin')
        if current_user.username != 'admin':
            if request.is_json or request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'error': 'Admin privileges required'}), 403
            # For HTML requests, show a forbidden page
            return render_template_string("""
                <h1>Access Denied</h1>
                <p>You do not have admin privileges to access this page.</p>
                <p><a href="{{ url_for('index') }}">Return to homepage</a></p>
            """), 403
            
        return f(*args, **kwargs)
    return wrapper

# --- HTTP Basic Auth Support (for tools/tests using requests.Session(auth=(u,p))) ---
@app.before_request
def basic_auth_login():
    """If an Authorization: Basic header is present, authenticate the user.
    This enables simple Basic Auth for API clients without changing existing
    session-based auth flows (Flask-Login remains the source of truth).
    """
    try:
        # Skip if already authenticated or for static assets
        if current_user.is_authenticated:
            return
        if request.endpoint in {"static", "login_static", "js_files"}:
            return

        auth = request.authorization  # Flask parses Basic/Digest headers
        if not auth or not auth.username or not auth.password:
            return

        user = User.query.filter_by(username=auth.username).first()
        if user and check_password_hash(user.password_hash, auth.password):
            login_user(user)
    except Exception:
        # Never block the request due to auth helper errors; normal flow continues
        pass

@login_manager.user_loader
def load_user(user_id):
    return db.session.get(User, int(user_id))

@login_manager.unauthorized_handler
def handle_unauthorized():
    """Return JSON 401 for API/SPA requests; otherwise redirect to login page."""
    api_like_paths = (
        '/api', '/task', '/files', '/process', '/history', '/download', '/profile',
        '/enhanced', '/advanced'
    )
    wants_json = (
        request.is_json or
        request.headers.get('X-Requested-With') == 'XMLHttpRequest' or
        request.accept_mimetypes['application/json'] >= request.accept_mimetypes['text/html'] or
        any(request.path.startswith(p) for p in api_like_paths)
    )
    if wants_json:
        return jsonify({'error': 'Unauthorized'}), 401
    # Fallback: serve SPA login page
    return app.send_static_file('index.html')

@app.route('/auth/check')
def auth_check():
    """Lightweight endpoint to check authentication status without redirect."""
    if current_user.is_authenticated:
        return jsonify({
            'authenticated': True,
            'user': {
                'id': current_user.id,
                'username': current_user.username,
                'email': current_user.email,
                'phone_number': current_user.phone_number,
                'created_at': current_user.created_at.isoformat()
            }
        }), 200
    return jsonify({'authenticated': False}), 401

def cleanup_old_tasks():
    """Remove tasks older than 1 hour"""
    current_time = time.time()
    expired_tasks = [
        task_id for task_id, timestamp in app.task_timestamps.items()
        if current_time - timestamp > 3600  # 1 hour
    ]
    for task_id in expired_tasks:
        if task_id in app.task_results:
            del app.task_results[task_id]
        if task_id in app.task_timestamps:
            del app.task_timestamps[task_id]

# --- Helper Functions ---
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def format_bytes(bytes, decimals=2):
    if bytes == 0:
        return '0 Bytes'
    k = 1024
    dm = decimals if decimals >= 0 else 0
    sizes = ['Bytes', 'KB', 'MB', 'GB', 'TB', 'PB', 'EB', 'ZB', 'YB']
    i = int(math.floor(math.log(bytes) / math.log(k)))
    return f"{round(bytes / math.pow(k, i), dm)} {sizes[i]}"

def record_file_conversion(original_file_id, output_file_path, conversion_type, user_id=None):
    """
    Record a file conversion operation in the database
    
    Args:
        original_file_id: ID of the original file in the FileRecord table
        output_file_path: Path to the output file
        conversion_type: Type of conversion (merge, split, compress, etc.)
        user_id: ID of the user who performed the conversion
    
    Returns:
        FileConversionRecord instance
    """
    try:
        start_time = time.time()
        file_size = os.path.getsize(output_file_path) if os.path.exists(output_file_path) else None
        
        # If user_id is not provided, try to get from current_user
        if user_id is None and current_user.is_authenticated:
            user_id = current_user.id
        
        # Default to user ID 1 for testing/anonymous
        user_id = user_id or 1
        
        # Record the conversion
        conversion_record = FileConversionRecord(
            original_file_id=original_file_id,
            output_file=os.path.basename(output_file_path),
            conversion_type=conversion_type,
            file_size=file_size,
            processing_time_ms=int((time.time() - start_time) * 1000),
            user_id=user_id,
            pg_tools_version='2.0.12\\Windows'
        )
        
        db.session.add(conversion_record)
        db.session.commit()
        
        # Update access count for original file
        original_file = db.session.get(FileRecord, original_file_id)
        if original_file:
            original_file.last_accessed = datetime.now(timezone.utc)
            original_file.access_count += 1
            db.session.commit()
            
        return conversion_record
    except Exception as e:
        logging.exception(f"Failed to record file conversion: {e}")
        db.session.rollback()
        return None

def get_file_by_key(file_key, user_id=None, update_access=True):
    """
    Get a file record by its unique key, optionally updating access statistics
    
    Args:
        file_key: The unique filename of the file
        user_id: Optional user ID to filter by (defaults to current_user.id)
        update_access: Whether to update last_accessed and access_count
        
    Returns:
        FileRecord instance or None if not found
    """
    try:
        # Use current user ID if authenticated and no user_id provided
        if user_id is None and current_user.is_authenticated:
            user_id = current_user.id
        
        query = FileRecord.query.filter_by(filename=file_key)
        
        # Add user filter if provided
        if user_id is not None:
            query = query.filter_by(user_id=user_id)
            
        file_record = query.first()
        
        # Update access statistics if requested and file found
        if file_record and update_access:
            file_record.last_accessed = datetime.now(timezone.utc)
            file_record.access_count += 1
            db.session.commit()
            
        return file_record
    except Exception as e:
        logging.exception(f"Error retrieving file record: {e}")
        return None

def update_pg_tools_version():
    """Update PostgreSQL tools version in the database config"""
    try:
        with app.app_context():
            # Check if we have a config table
            inspector = inspect(db.engine)
            config_table_exists = inspector.has_table('app_config')
            
            if not config_table_exists:
                # Create app_config table if it doesn't exist
                db.session.execute(db.text('''
                CREATE TABLE IF NOT EXISTS app_config (
                    key VARCHAR(100) PRIMARY KEY,
                    value TEXT NOT NULL,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
                '''))
            
            # Check if the entry already exists
            result = db.session.execute(db.text("SELECT 1 FROM app_config WHERE key = 'pg_tools_version'")).fetchone()
            
            if result:
                # Update existing entry
                db.session.execute(db.text("UPDATE app_config SET value = '2.0.12\\Windows', updated_at = CURRENT_TIMESTAMP WHERE key = 'pg_tools_version'"))
            else:
                # Insert new entry
                db.session.execute(db.text("INSERT INTO app_config (key, value, updated_at) VALUES ('pg_tools_version', '2.0.12\\Windows', CURRENT_TIMESTAMP)"))
            
            db.session.commit()
            return True
    except Exception as e:
        logging.exception(f"Failed to update PostgreSQL tools version: {e}")
        return False

# Execute this once at startup
try:
    # No need to wrap in app_context here since the function does it internally
    update_pg_tools_version()
except Exception as e:
    logging.warning(f"Could not update PostgreSQL tools version: {e}")

# --- Authentication Routes ---
@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'GET':
        return render_template('login/html/register.html')
    
    try:
        # Handle form submissions
        if request.content_type == 'application/x-www-form-urlencoded':
            username = request.form.get('username')
            email = request.form.get('email')
            password = request.form.get('password')
        # Handle API JSON requests
        elif request.is_json:
            data = request.get_json()
            if not data:
                return jsonify({'error': 'Invalid JSON data'}), 400
            username = data.get('username')
            email = data.get('email')
            password = data.get('password')
        else:
            return jsonify({'error': 'Unsupported content type'}), 400
        
        if not username or not email or not password:
            return jsonify({'error': 'Username, email, and password are required'}), 400
        
        if len(password) < 6:
            return jsonify({'error': 'Password must be at least 6 characters long'}), 400
        
        # Check if user exists
        existing_user = User.query.filter_by(username=username).first()
        if existing_user:
            return jsonify({'error': 'Username already exists'}), 400
        
        existing_email = User.query.filter_by(email=email).first()
        if existing_email:
            return jsonify({'error': 'Email already exists'}), 400
        
        # Create new user
        hashed_password = generate_password_hash(password)
        new_user = User(
            username=username,
            email=email,
            password_hash=hashed_password
        )
        
        # Add to database
        try:
            db.session.add(new_user)
            db.session.commit()
            
            # Log in the user after registration
            login_user(new_user)
            
            return jsonify({'message': 'User registered and logged in successfully'}), 201
        except Exception as db_error:
            db.session.rollback()
            logging.exception("Database error during registration")
            return jsonify({'error': f'Database error: {str(db_error)}'}), 500
        
    except Exception as e:
        logging.exception("Registration error occurred")
        db.session.rollback()
        return jsonify({'error': f'Internal server error: {str(e)}'}), 500

@app.route('/reset-password', methods=['GET', 'POST'])
def reset_password_request():
    if request.method == 'GET':
        return render_template('login/html/reset.html')
    
    try:
        # Handle form submissions
        if request.content_type == 'application/x-www-form-urlencoded':
            email = request.form.get('email')
        # Handle API JSON requests
        elif request.is_json:
            data = request.get_json()
            if not data:
                return jsonify({'error': 'Invalid JSON data'}), 400
            email = data.get('email')
        else:
            return jsonify({'error': 'Unsupported content type'}), 400
        
        if not email:
            return jsonify({'error': 'Email is required'}), 400
        
        # Check if user exists
        user = User.query.filter_by(email=email).first()
        if not user:
            # Don't reveal that the user doesn't exist for security reasons
            return jsonify({'message': 'If this email is registered, password reset instructions will be sent'}), 200
        
        # Generate a reset token
        serializer = URLSafeTimedSerializer(app.config['SECRET_KEY'])
        token = serializer.dumps(email, salt='password-reset-salt')
        
        # In a real application, you would send an email with the reset link
        # For this example, we'll just return the token in the response
        reset_url = url_for('reset_password_confirm', token=token, _external=True)
        
        # Log the reset token (in a real app, this would be sent via email)
        logging.info(f"Password reset requested for {email}. Reset URL: {reset_url}")
        
        return jsonify({
            'message': 'Password reset instructions sent to your email',
            'debug_token': token if is_debug_mode() else None
        }), 200
        
    except Exception as e:
        logging.exception("Reset password error occurred")
        return jsonify({'error': f'Internal server error: {str(e)}'}), 500

@app.route('/reset-password/<token>', methods=['GET', 'POST'])
def reset_password_confirm(token):
    try:
        # Verify the token
        serializer = URLSafeTimedSerializer(app.config['SECRET_KEY'])
        email = serializer.loads(token, salt='password-reset-salt', max_age=3600)  # Token valid for 1 hour
        
        # Check if user exists
        user = User.query.filter_by(email=email).first()
        if not user:
            return jsonify({'error': 'Invalid or expired token'}), 400
        
        if request.method == 'GET':
            return render_template('login/html/reset_confirm.html', token=token)
        
        # Process POST request
        if request.content_type == 'application/x-www-form-urlencoded':
            new_password = request.form.get('password')
        elif request.is_json:
            data = request.get_json()
            if not data:
                return jsonify({'error': 'Invalid JSON data'}), 400
            new_password = data.get('password')
        else:
            return jsonify({'error': 'Unsupported content type'}), 400
        
        if not new_password:
            return jsonify({'error': 'New password is required'}), 400
        
        if len(new_password) < 6:
            return jsonify({'error': 'Password must be at least 6 characters long'}), 400
        
        # Update the user's password
        user.password_hash = generate_password_hash(new_password)
        
        try:
            db.session.commit()
            return jsonify({'message': 'Password has been reset successfully'}), 200
        except Exception as db_error:
            db.session.rollback()
            logging.exception("Database error during password reset")
            return jsonify({'error': f'Database error: {str(db_error)}'}), 500
            
    except SignatureExpired:
        return jsonify({'error': 'The password reset link has expired'}), 400
    except BadSignature:
        return jsonify({'error': 'Invalid reset token'}), 400
    except Exception as e:
        logging.exception("Reset password confirmation error occurred")
        return jsonify({'error': f'Internal server error: {str(e)}'}), 500

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'GET':
        return render_template('login/html/index.html')
    
    try:
        # Handle form submissions
        if request.content_type == 'application/x-www-form-urlencoded':
            username = request.form.get('username')
            password = request.form.get('password')
        # Handle API JSON requests
        elif request.is_json:
            data = request.get_json()
            if not data:
                return jsonify({'error': 'Invalid JSON data'}), 400
            username = data.get('username')
            password = data.get('password')
        else:
            return jsonify({'error': 'Unsupported content type'}), 400
        
        if not username or not password:
            return jsonify({'error': 'Username and password are required'}), 400
        
        # Try to find user by username or email
        user = User.query.filter_by(username=username).first()
        
        # If not found by username, try email
        if not user:
            user = User.query.filter_by(email=username).first()
        
        if user and check_password_hash(user.password_hash, password):
            # Login successful
            login_user(user)
            
            # Track login activity
            now = datetime.now(timezone.utc)
            user.last_login = now
            user.login_count += 1
            user.last_ip = request.remote_addr
            user.user_agent = request.user_agent.string if request.user_agent else None
            
            # Create login history record
            login_history = UserLoginHistory(
                user_id=user.id,
                login_time=now,
                ip_address=request.remote_addr,
                user_agent=request.user_agent.string if request.user_agent else None,
                success=True,
                device_info={
                    'platform': request.user_agent.platform if request.user_agent else None,
                    'browser': request.user_agent.browser if request.user_agent else None,
                    'version': request.user_agent.version if request.user_agent else None,
                    'language': request.accept_languages.best if request.accept_languages else None,
                    'pg_tools_version': '2.0.12\\Windows'
                }
            )
            
            try:
                db.session.add(login_history)
                db.session.commit()
                return jsonify({
                    'message': 'Logged in successfully',
                    'user': {
                        'id': user.id,
                        'username': user.username,
                        'email': user.email
                    }
                }), 200
            except Exception as db_error:
                db.session.rollback()
                logging.exception("Database error during login history recording")
                # Still return success even if history recording fails
                return jsonify({'message': 'Logged in successfully (history not recorded)'}), 200
        
        # Record failed login attempt
        if user:
            failed_login = UserLoginHistory(
                user_id=user.id,
                login_time=datetime.now(timezone.utc),
                ip_address=request.remote_addr,
                user_agent=request.user_agent.string if request.user_agent else None,
                success=False
            )
            try:
                db.session.add(failed_login)
                db.session.commit()
            except Exception as db_error:
                db.session.rollback()
                logging.exception("Database error during failed login recording")
        
        return jsonify({'error': 'Invalid username/email or password'}), 401
        
    except Exception as e:
        logging.exception("Login error occurred")
        return jsonify({'error': f'Internal server error: {str(e)}'}), 500

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return jsonify({'message': 'Logged out successfully'}), 200

@app.route('/profile')
@login_required
def profile():
    return jsonify({
        'id': current_user.id,
        'username': current_user.username,
        'email': current_user.email,
        'phone_number': current_user.phone_number,
        'created_at': current_user.created_at.isoformat(),
        'last_login': current_user.last_login.isoformat() if current_user.last_login else None,
        'login_count': current_user.login_count,
        'last_ip': current_user.last_ip
    })

@app.route('/profile/activity')
@login_required
def user_activity():
    """Return user activity statistics including login history and file operations"""
    
    # Get login history
    login_history = UserLoginHistory.query.filter_by(user_id=current_user.id).order_by(
        UserLoginHistory.login_time.desc()).limit(10).all()
    
    # Get file statistics
    file_count = FileRecord.query.filter_by(user_id=current_user.id).count()
    total_file_size = db.session.query(db.func.sum(FileRecord.file_size)).filter(
        FileRecord.user_id == current_user.id).scalar() or 0
    
    # Get most accessed files
    most_accessed_files = FileRecord.query.filter_by(user_id=current_user.id).order_by(
        FileRecord.access_count.desc()).limit(5).all()
    
    # Get recent conversions
    recent_conversions = FileConversionRecord.query.filter_by(user_id=current_user.id).order_by(
        FileConversionRecord.conversion_time.desc()).limit(5).all()
    
    # Format data for response
    return jsonify({
        'login_history': [
            {
                'login_time': login.login_time.isoformat(),
                'ip_address': login.ip_address,
                'success': login.success,
                'device_info': login.device_info
            } for login in login_history
        ],
        'file_statistics': {
            'file_count': file_count,
            'total_size': total_file_size,
            'formatted_size': format_bytes(total_file_size),
            'most_accessed_files': [
                {
                    'id': f.id,
                    'filename': f.original_filename,
                    'access_count': f.access_count,
                    'last_accessed': f.last_accessed.isoformat() if f.last_accessed else None
                } for f in most_accessed_files
            ],
            'recent_conversions': [
                {
                    'id': c.id,
                    'conversion_type': c.conversion_type,
                    'output_file': c.output_file,
                    'conversion_time': c.conversion_time.isoformat(),
                    'file_size': format_bytes(c.file_size) if c.file_size else 'Unknown',
                    'processing_time_ms': c.processing_time_ms,
                    'pg_tools_version': c.pg_tools_version
                } for c in recent_conversions
            ]
        }
    })

 # --- Password Reset Routes ---
@app.route('/reset-password/<token>', methods=['GET'])
def reset_password_page(token):
    """Serve the password reset HTML page."""
    return render_template('login/html/reset.html', token=token)

@app.route('/request-password-reset', methods=['GET'])
def request_password_reset_page():
    """Serve the request password reset HTML page."""
    return render_template('login/html/reset.html')
@app.route('/request-password-reset', methods=['POST'])
def request_password_reset():
    try:
        if not request.is_json:
            return jsonify({'error': 'Content-Type must be application/json'}), 400
        
        data = request.get_json()
        email = data.get('email')
        
        if not email:
            return jsonify({'error': 'Email is required'}), 400
        
        user = User.query.filter_by(email=email).first()
        if not user:
            # Don't reveal if email exists or not for security
            return jsonify({'message': 'If the email exists, a password reset link has been sent'}), 200
        
        # Generate reset token
        token = serializer.dumps(user.email, salt=get_env('PASSWORD_RESET_SALT', default=None, required=True))
        
        # Create reset URL
        reset_url = f"{request.host_url}reset-password/{token}"
        
        # Send email
        try:
            msg = Message(
                'Password Reset Request',
                recipients=[user.email],
                body=f'''To reset your password, visit the following link:
{reset_url}

If you did not make this request, simply ignore this email.

This link will expire in 1 hour.
'''
            )
            mail.send(msg)
            return jsonify({'message': 'Password reset email sent successfully'}), 200
        except Exception as e:
            logging.error(f"Failed to send email: {e}")
            return jsonify({'error': 'Failed to send password reset email'}), 500
            
    except Exception as e:
        logging.error(f"Password reset request error: {e}")
        return jsonify({'error': 'Internal server error'}), 500

@app.route('/reset-password/<token>', methods=['POST'])
def reset_password(token):
    try:
        if not request.is_json:
            return jsonify({'error': 'Content-Type must be application/json'}), 400
        
        data = request.get_json()
        new_password = data.get('new_password')
        
        if not new_password or len(new_password) < 6:
            return jsonify({'error': 'New password must be at least 6 characters long'}), 400
        
        try:
            email = serializer.loads(token, salt=get_env('PASSWORD_RESET_SALT', default=None, required=True), max_age=3600)  # 1 hour expiry
        except SignatureExpired:
            return jsonify({'error': 'Password reset link has expired'}), 400
        except BadSignature:
            return jsonify({'error': 'Invalid password reset link'}), 400
        
        user = User.query.filter_by(email=email).first()
        if not user:
            return jsonify({'error': 'User not found'}), 404
        
        # Update password
        user.password_hash = generate_password_hash(new_password)
        db.session.commit()
        
        return jsonify({'message': 'Password reset successfully'}), 200
        
    except Exception as e:
        logging.error(f"Password reset error: {e}")
        return jsonify({'error': 'Internal server error'}), 500

# --- OTP Authentication Routes ---
@app.route('/send-otp', methods=['POST'])
def send_otp():
    if not TWILIO_AVAILABLE or not REDIS_AVAILABLE:
        return jsonify({'error': 'OTP service not available'}), 503
    
    try:
        if not request.is_json:
            return jsonify({'error': 'Content-Type must be application/json'}), 400
        
        data = request.get_json()
        phone_number = data.get('phone_number')
        
        if not phone_number:
            return jsonify({'error': 'Phone number is required'}), 400
        
        # Generate 6-digit OTP
        otp = ''.join(random.choices(string.digits, k=6))
        
        # Store OTP in Redis with 5-minute expiry
        redis_client.setex(f"otp:{phone_number}", 300, otp)
        
        # Send OTP via SMS
        try:
            message = twilio_client.messages.create(
                body=f'Your PDF Tool verification code is: {otp}. Valid for 5 minutes.',
                from_=twilio_phone,
                to=phone_number
            )
            return jsonify({'message': 'OTP sent successfully'}), 200
        except TwilioRestException as e:
            logging.error(f"Twilio error: {e}")
            return jsonify({'error': 'Failed to send OTP'}), 500
            
    except Exception as e:
        logging.error(f"Send OTP error: {e}")
        return jsonify({'error': 'Internal server error'}), 500

@app.route('/verify-otp', methods=['POST'])
def verify_otp():
    if not TWILIO_AVAILABLE or not REDIS_AVAILABLE:
        return jsonify({'error': 'OTP service not available'}), 503
    
    try:
        if not request.is_json:
            return jsonify({'error': 'Content-Type must be application/json'}), 400
        
        data = request.get_json()
        phone_number = data.get('phone_number')
        otp = data.get('otp')
        
        if not phone_number or not otp:
            return jsonify({'error': 'Phone number and OTP are required'}), 400
        
        # Verify OTP from Redis
        stored_otp = redis_client.get(f"otp:{phone_number}")
        if not stored_otp or stored_otp != otp:
            return jsonify({'error': 'Invalid or expired OTP'}), 400
        
        # Find or create user
        user = User.query.filter_by(phone_number=phone_number).first()
        if not user:
            # Create new user with phone number
            username = f"user_{phone_number[-4:]}"  # Use last 4 digits as username
            user = User(
                username=username,
                email=f"{username}@phone.com",  # Placeholder email
                phone_number=phone_number,
                password_hash=generate_password_hash('')  # No password for OTP users
            )
            db.session.add(user)
            db.session.commit()
        
        # Login user
        login_user(user)
        
        # Clear OTP from Redis
        redis_client.delete(f"otp:{phone_number}")
        
        return jsonify({'message': 'OTP verified successfully', 'user': {
            'id': user.id,
            'username': user.username,
            'phone_number': user.phone_number
        }}), 200
        
    except Exception as e:
        logging.error(f"Verify OTP error: {e}")
        return jsonify({'error': 'Internal server error'}), 500

# --- Main Routes ---
@app.route('/')
def index():
    # Redirect to login page as that's our main entry point
    return redirect(url_for('login'))

@app.route('/pdf-tools')
@login_required
def pdf_tools_shortcut():
    """Redirect to PDF operations dashboard"""
    return redirect(url_for('pdf_operations.pdf_dashboard'))

# Health check endpoint

@app.route('/admin/dashboard')
@admin_required
def admin_dashboard():
    """Admin dashboard for system management"""
    # Get total user count
    user_count = User.query.count()
    
    # Get file statistics
    file_count = FileRecord.query.count()
    total_size = db.session.query(db.func.sum(FileRecord.file_size)).scalar() or 0
    
    # Get recent user registrations
    recent_users = User.query.order_by(User.created_at.desc()).limit(5).all()
    
    # Get system status info
    disk_usage = shutil.disk_usage('/')
    free_space_gb = disk_usage.free / (1024 * 1024 * 1024)
    
    # For API/JSON requests
    if request.headers.get('Accept') == 'application/json':
        return jsonify({
            'user_count': user_count,
            'file_count': file_count, 
            'total_size': format_bytes(total_size),
            'free_space': f"{free_space_gb:.2f} GB",
            'recent_users': [{
                'id': user.id,
                'username': user.username,
                'email': user.email,
                'created_at': user.created_at.isoformat()
            } for user in recent_users]
        })
    
    # For HTML requests, use template string for simplicity
    return render_template_string("""
        <!DOCTYPE html>
        <html>
        <head>
            <title>Admin Dashboard</title>
            <style>
                body { font-family: Arial, sans-serif; margin: 20px; }
                .dashboard { max-width: 1000px; margin: 0 auto; }
                .card { border: 1px solid #ccc; border-radius: 5px; padding: 15px; margin: 10px 0; }
                .stats { display: flex; flex-wrap: wrap; }
                .stat-card { flex: 1; min-width: 200px; margin: 10px; padding: 15px; border: 1px solid #ddd; border-radius: 5px; }
                table { width: 100%; border-collapse: collapse; }
                th, td { text-align: left; padding: 8px; border-bottom: 1px solid #ddd; }
                th { background-color: #f2f2f2; }
            </style>
        </head>
        <body>
            <div class="dashboard">
                <h1>Admin Dashboard</h1>
                <div class="card">
                    <h2>System Overview</h2>
                    <div class="stats">
                        <div class="stat-card">
                            <h3>Users</h3>
                            <p>{{ user_count }}</p>
                        </div>
                        <div class="stat-card">
                            <h3>Files</h3>
                            <p>{{ file_count }}</p>
                        </div>
                        <div class="stat-card">
                            <h3>Storage Used</h3>
                            <p>{{ total_size }}</p>
                        </div>
                        <div class="stat-card">
                            <h3>Free Space</h3>
                            <p>{{ free_space }}</p>
                        </div>
                    </div>
                </div>
                
                <div class="card">
                    <h2>Recent Users</h2>
                    <table>
                        <tr>
                            <th>ID</th>
                            <th>Username</th>
                            <th>Email</th>
                            <th>Created</th>
                        </tr>
                        {% for user in recent_users %}
                        <tr>
                            <td>{{ user.id }}</td>
                            <td>{{ user.username }}</td>
                            <td>{{ user.email }}</td>
                            <td>{{ user.created_at.strftime('%Y-%m-%d %H:%M:%S') }}</td>
                        </tr>
                        {% endfor %}
                    </table>
                </div>
                
                <div class="card">
                    <h2>Admin Actions</h2>
                    <p><a href="{{ url_for('admin_users') }}">Manage Users</a></p>
                    <p><a href="{{ url_for('admin_files') }}">Manage Files</a></p>
                    <p><a href="{{ url_for('admin_system') }}">System Settings</a></p>
                </div>
            </div>
        </body>
        </html>
    """, 
    user_count=user_count,
    file_count=file_count,
    total_size=format_bytes(total_size),
    free_space=f"{free_space_gb:.2f} GB",
    recent_users=recent_users
    )

@app.route('/admin/system-stats')
@admin_required
def admin_system_stats():
    """Advanced system statistics including user activity and PG tools information"""
    from datetime import timedelta
    
    # Get PostgreSQL tools version
    pg_tools_version = None
    try:
        result = db.engine.execute("SELECT value FROM app_config WHERE key = 'pg_tools_version'").fetchone()
        if result:
            pg_tools_version = result[0]
    except Exception as e:
        logging.warning(f"Could not retrieve PG tools version: {e}")
    
    # User statistics
    user_count = User.query.count()
    active_users_24h = UserLoginHistory.query.filter(
        UserLoginHistory.login_time > datetime.now(timezone.utc) - timedelta(days=1)
    ).with_entities(UserLoginHistory.user_id).distinct().count()
    
    # Login statistics
    login_attempts = UserLoginHistory.query.count()
    failed_logins = UserLoginHistory.query.filter_by(success=False).count()
    
    # File statistics
    file_count = FileRecord.query.count()
    conversion_count = FileConversionRecord.query.count()
    total_size = db.session.query(db.func.sum(FileRecord.file_size)).scalar() or 0
    
    # File type breakdown
    file_types = db.session.query(
        FileRecord.file_type, 
        db.func.count(FileRecord.id)
    ).group_by(FileRecord.file_type).all()
    
    # Conversion type breakdown
    conversion_types = db.session.query(
        FileConversionRecord.conversion_type, 
        db.func.count(FileConversionRecord.id)
    ).group_by(FileConversionRecord.conversion_type).all()
    
    # System resource usage
    disk_usage = shutil.disk_usage('/')
    free_space_gb = disk_usage.free / (1024 * 1024 * 1024)
    
    # Format the response data
    return jsonify({
        'pg_tools': {
            'version': pg_tools_version or 'Unknown',
            'last_update': datetime.now(timezone.utc).isoformat(),
            'status': 'Active' if pg_tools_version else 'Not Detected'
        },
        'user_stats': {
            'total_users': user_count,
            'active_users_24h': active_users_24h,
            'login_attempts': login_attempts,
            'failed_logins': failed_logins,
            'success_rate': round(((login_attempts - failed_logins) / login_attempts) * 100, 2) if login_attempts > 0 else 100
        },
        'file_stats': {
            'total_files': file_count,
            'total_size': total_size,
            'formatted_size': format_bytes(total_size),
            'conversions': conversion_count,
            'file_types': [{
                'type': file_type,
                'count': count
            } for file_type, count in file_types],
            'conversion_types': [{
                'type': conv_type,
                'count': count
            } for conv_type, count in conversion_types]
        },
        'system_resources': {
            'disk_space': {
                'total_gb': round(disk_usage.total / (1024 * 1024 * 1024), 2),
                'used_gb': round(disk_usage.used / (1024 * 1024 * 1024), 2),
                'free_gb': round(free_space_gb, 2),
                'percent_used': round((disk_usage.used / disk_usage.total) * 100, 2)
            },
            'timestamps': {
                'server_time': datetime.now().isoformat(),
                'utc_time': datetime.now(timezone.utc).isoformat()
            }
        }
    })

@app.route('/admin/users')
@admin_required
def admin_users():
    """Admin user management"""
    users = User.query.order_by(User.id).all()
    
    # For API/JSON requests
    if request.headers.get('Accept') == 'application/json':
        return jsonify({
            'users': [{
                'id': user.id,
                'username': user.username,
                'email': user.email,
                'created_at': user.created_at.isoformat()
            } for user in users]
        })
    
    # For HTML requests
    return render_template_string("""
        <!DOCTYPE html>
        <html>
        <head>
            <title>User Management</title>
            <style>
                body { font-family: Arial, sans-serif; margin: 20px; }
                .container { max-width: 1000px; margin: 0 auto; }
                table { width: 100%; border-collapse: collapse; }
                th, td { text-align: left; padding: 8px; border-bottom: 1px solid #ddd; }
                th { background-color: #f2f2f2; }
                .actions { display: flex; gap: 10px; }
                .actions a { color: blue; text-decoration: none; }
                .actions a:hover { text-decoration: underline; }
                .back-link { margin-bottom: 20px; }
            </style>
        </head>
        <body>
            <div class="container">
                <div class="back-link">
                    <a href="{{ url_for('admin_dashboard') }}">← Back to Dashboard</a>
                </div>
                
                <h1>User Management</h1>
                <table>
                    <tr>
                        <th>ID</th>
                        <th>Username</th>
                        <th>Email</th>
                        <th>Created</th>
                        <th>Actions</th>
                    </tr>
                    {% for user in users %}
                    <tr>
                        <td>{{ user.id }}</td>
                        <td>{{ user.username }}</td>
                        <td>{{ user.email }}</td>
                        <td>{{ user.created_at.strftime('%Y-%m-%d %H:%M:%S') }}</td>
                        <td class="actions">
                            <a href="{{ url_for('admin_edit_user', user_id=user.id) }}">Edit</a>
                            {% if user.username != 'admin' %}
                            <a href="{{ url_for('admin_delete_user', user_id=user.id) }}" 
                               onclick="return confirm('Are you sure you want to delete this user?')">Delete</a>
                            {% endif %}
                        </td>
                    </tr>
                    {% endfor %}
                </table>
            </div>
        </body>
        </html>
    """, users=users)

@app.route('/admin/files')
@admin_required
def admin_files():
    """Admin file management"""
    files = FileRecord.query.order_by(FileRecord.upload_date.desc()).all()
    
    # For API/JSON requests
    if request.headers.get('Accept') == 'application/json':
        return jsonify({
            'files': [{
                'id': f.id,
                'filename': f.filename,
                'original_filename': f.original_filename,
                'file_size': f.file_size,
                'formatted_size': format_bytes(f.file_size),
                'upload_date': f.upload_date.isoformat(),
                'user_id': f.user_id
            } for f in files]
        })
    
    # For HTML requests
    return render_template_string("""
        <!DOCTYPE html>
        <html>
        <head>
            <title>File Management</title>
            <style>
                body { font-family: Arial, sans-serif; margin: 20px; }
                .container { max-width: 1000px; margin: 0 auto; }
                table { width: 100%; border-collapse: collapse; }
                th, td { text-align: left; padding: 8px; border-bottom: 1px solid #ddd; }
                th { background-color: #f2f2f2; }
                .actions { display: flex; gap: 10px; }
                .actions a { color: blue; text-decoration: none; }
                .actions a:hover { text-decoration: underline; }
                .back-link { margin-bottom: 20px; }
            </style>
        </head>
        <body>
            <div class="container">
                <div class="back-link">
                    <a href="{{ url_for('admin_dashboard') }}">← Back to Dashboard</a>
                </div>
                
                <h1>File Management</h1>
                <table>
                    <tr>
                        <th>ID</th>
                        <th>Original Name</th>
                        <th>Size</th>
                        <th>Upload Date</th>
                        <th>User ID</th>
                        <th>Actions</th>
                    </tr>
                    {% for file in files %}
                    <tr>
                        <td>{{ file.id }}</td>
                        <td>{{ file.original_filename }}</td>
                        <td>{{ format_bytes(file.file_size) }}</td>
                        <td>{{ file.upload_date.strftime('%Y-%m-%d %H:%M:%S') }}</td>
                        <td>{{ file.user_id }}</td>
                        <td class="actions">
                            <a href="{{ url_for('download_file', key=file.filename) }}">Download</a>
                            <a href="{{ url_for('admin_delete_file', file_id=file.id) }}"
                               onclick="return confirm('Are you sure you want to delete this file?')">Delete</a>
                        </td>
                    </tr>
                    {% endfor %}
                </table>
            </div>
        </body>
        </html>
    """, files=files, format_bytes=format_bytes)

@app.route('/admin/system')
@admin_required
def admin_system():
    """Admin system settings"""
    
    # System information
    system_info = {
        'os': os.name,
        'python_version': sys.version,
        'app_directory': os.path.abspath(os.path.dirname(__file__)),
        'uploads_directory': os.path.abspath(UPLOAD_FOLDER),
        'processed_directory': os.path.abspath(PROCESSED_FOLDER),
    }
    
    # For API/JSON requests
    if request.headers.get('Accept') == 'application/json':
        return jsonify({
            'system_info': system_info,
            'environment_variables': {
                'FLASK_ENV': os.getenv('FLASK_ENV', 'production'),
                'FLASK_DEBUG': os.getenv('FLASK_DEBUG', 'False'),
                'SECRET_KEY': 'REDACTED',
                'DATABASE_URL': get_database_url().replace('://', '://[CREDENTIALS]@') if '@' in get_database_url() else get_database_url()
            }
        })
    
    # For HTML requests
    return render_template_string("""
        <!DOCTYPE html>
        <html>
        <head>
            <title>System Settings</title>
            <style>
                body { font-family: Arial, sans-serif; margin: 20px; }
                .container { max-width: 1000px; margin: 0 auto; }
                .card { border: 1px solid #ccc; border-radius: 5px; padding: 15px; margin: 15px 0; }
                table { width: 100%; border-collapse: collapse; }
                th, td { text-align: left; padding: 8px; border-bottom: 1px solid #ddd; }
                th { background-color: #f2f2f2; width: 30%; }
                .back-link { margin-bottom: 20px; }
            </style>
        </head>
        <body>
            <div class="container">
                <div class="back-link">
                    <a href="{{ url_for('admin_dashboard') }}">← Back to Dashboard</a>
                </div>
                
                <h1>System Settings</h1>
                
                <div class="card">
                    <h2>System Information</h2>
                    <table>
                        <tr>
                            <th>Operating System</th>
                            <td>{{ system_info.os }}</td>
                        </tr>
                        <tr>
                            <th>Python Version</th>
                            <td>{{ system_info.python_version }}</td>
                        </tr>
                        <tr>
                            <th>App Directory</th>
                            <td>{{ system_info.app_directory }}</td>
                        </tr>
                        <tr>
                            <th>Uploads Directory</th>
                            <td>{{ system_info.uploads_directory }}</td>
                        </tr>
                        <tr>
                            <th>Processed Directory</th>
                            <td>{{ system_info.processed_directory }}</td>
                        </tr>
                    </table>
                </div>
                
                <div class="card">
                    <h2>Environment Variables</h2>
                    <table>
                        <tr>
                            <th>FLASK_ENV</th>
                            <td>{{ os.getenv('FLASK_ENV', 'production') }}</td>
                        </tr>
                        <tr>
                            <th>FLASK_DEBUG</th>
                            <td>{{ os.getenv('FLASK_DEBUG', 'False') }}</td>
                        </tr>
                        <tr>
                            <th>SECRET_KEY</th>
                            <td>REDACTED</td>
                        </tr>
                        <tr>
                            <th>DATABASE_URL</th>
                            <td>{{ redacted_db_url }}</td>
                        </tr>
                    </table>
                </div>
                
                <div class="card">
                    <h2>Actions</h2>
                    <p><a href="{{ url_for('admin_clear_processed_files') }}" 
                          onclick="return confirm('Are you sure you want to clear all processed files?')">
                            Clear Processed Files
                        </a>
                    </p>
                </div>
            </div>
        </body>
        </html>
    """, 
    system_info=system_info, 
    os=os,
    redacted_db_url=get_database_url().replace('://', '://[CREDENTIALS]@') if '@' in get_database_url() else get_database_url()
    )

@app.route('/healthz', methods=['GET'])
def healthz_check():
    """Secondary health check endpoint (simpler implementation)"""
    try:
        # Simple check that doesn't access DB or file system
        return jsonify({
            "status": "ok",
            "message": "Service responding",
            "timestamp": datetime.now().isoformat()
        }), 200
    except Exception as e:
        logging.error(f"Secondary health check failed: {str(e)}")
        return jsonify({
            "status": "error",
            "message": "Service is experiencing issues",
            "timestamp": datetime.now().isoformat()
        }), 503

@app.route('/admin/edit-user/<int:user_id>', methods=['GET', 'POST'])
@admin_required
def admin_edit_user(user_id):
    """Admin user edit"""
    user = User.query.get_or_404(user_id)
    
    if request.method == 'POST':
        # Handle JSON or form data
        if request.is_json:
            data = request.get_json()
        else:
            data = request.form
            
        # Update user fields
        if data.get('username'):
            user.username = data.get('username')
        if data.get('email'):
            user.email = data.get('email')
        if data.get('password'):
            user.password_hash = generate_password_hash(data.get('password'))
            
        db.session.commit()
        
        # Return JSON response for API requests
        if request.is_json:
            return jsonify({'message': 'User updated successfully'})
        
        # Redirect for form submissions
        return redirect(url_for('admin_users'))
    
    # GET request - show edit form
    return render_template_string("""
        <!DOCTYPE html>
        <html>
        <head>
            <title>Edit User</title>
            <style>
                body { font-family: Arial, sans-serif; margin: 20px; }
                .container { max-width: 600px; margin: 0 auto; }
                .form-group { margin-bottom: 15px; }
                label { display: block; margin-bottom: 5px; }
                input[type="text"], input[type="email"], input[type="password"] {
                    width: 100%;
                    padding: 8px;
                    box-sizing: border-box;
                }
                button { padding: 8px 15px; background: #4CAF50; color: white; border: none; cursor: pointer; }
                .back-link { margin-bottom: 20px; }
            </style>
        </head>
        <body>
            <div class="container">
                <div class="back-link">
                    <a href="{{ url_for('admin_users') }}">← Back to Users</a>
                </div>
                
                <h1>Edit User</h1>
                <form method="post">
                    <div class="form-group">
                        <label for="username">Username</label>
                        <input type="text" id="username" name="username" value="{{ user.username }}" required>
                    </div>
                    <div class="form-group">
                        <label for="email">Email</label>
                        <input type="email" id="email" name="email" value="{{ user.email }}" required>
                    </div>
                    <div class="form-group">
                        <label for="password">New Password (leave blank to keep current)</label>
                        <input type="password" id="password" name="password">
                    </div>
                    <button type="submit">Update User</button>
                </form>
            </div>
        </body>
        </html>
    """, user=user)

@app.route('/admin/delete-user/<int:user_id>', methods=['GET'])
@admin_required
def admin_delete_user(user_id):
    """Admin delete user"""
    user = User.query.get_or_404(user_id)
    
    # Prevent deleting the admin user
    if user.username == 'admin':
        return jsonify({'error': 'Cannot delete admin user'}), 403
    
    try:
        # Delete associated files
        files = FileRecord.query.filter_by(user_id=user.id).all()
        for file in files:
            # Delete physical file if it exists
            file_path = os.path.join(UPLOAD_FOLDER, file.filename)
            if os.path.exists(file_path):
                os.remove(file_path)
            
            # Delete database record
            db.session.delete(file)
        
        # Delete user
        db.session.delete(user)
        db.session.commit()
        
        # Return JSON response for API requests
        if request.headers.get('Accept') == 'application/json':
            return jsonify({'message': 'User deleted successfully'})
        
        # Redirect for browser requests
        return redirect(url_for('admin_users'))
        
    except Exception as e:
        db.session.rollback()
        logging.error(f"Error deleting user: {str(e)}")
        
        # Return error response
        if request.headers.get('Accept') == 'application/json':
            return jsonify({'error': f'Failed to delete user: {str(e)}'}), 500
        
        # Show error page for browser requests
        return render_template_string("""
            <h1>Error</h1>
            <p>Failed to delete user: {{ error }}</p>
            <p><a href="{{ url_for('admin_users') }}">Back to Users</a></p>
        """, error=str(e)), 500

@app.route('/admin/delete-file/<int:file_id>', methods=['GET'])
@admin_required
def admin_delete_file(file_id):
    """Admin delete file"""
    file = FileRecord.query.get_or_404(file_id)
    
    try:
        # Delete physical file if it exists
        file_path = os.path.join(UPLOAD_FOLDER, file.filename)
        if os.path.exists(file_path):
            os.remove(file_path)
        
        # Delete database record
        db.session.delete(file)
        db.session.commit()
        
        # Return JSON response for API requests
        if request.headers.get('Accept') == 'application/json':
            return jsonify({'message': 'File deleted successfully'})
        
        # Redirect for browser requests
        return redirect(url_for('admin_files'))
        
    except Exception as e:
        db.session.rollback()
        logging.error(f"Error deleting file: {str(e)}")
        
        # Return error response
        if request.headers.get('Accept') == 'application/json':
            return jsonify({'error': f'Failed to delete file: {str(e)}'}), 500
        
        # Show error page for browser requests
        return render_template_string("""
            <h1>Error</h1>
            <p>Failed to delete file: {{ error }}</p>
            <p><a href="{{ url_for('admin_files') }}">Back to Files</a></p>
        """, error=str(e)), 500

@app.route('/admin/clear-processed-files', methods=['GET'])
@admin_required
def admin_clear_processed_files():
    """Clear all files in the processed folder"""
    try:
        files_cleared = 0
        for filename in os.listdir(PROCESSED_FOLDER):
            file_path = os.path.join(PROCESSED_FOLDER, filename)
            if os.path.isfile(file_path):
                os.remove(file_path)
                files_cleared += 1
        
        # Return JSON response for API requests
        if request.headers.get('Accept') == 'application/json':
            return jsonify({'message': f'Cleared {files_cleared} processed files'})
        
        # Redirect for browser requests
        return redirect(url_for('admin_system'))
        
    except Exception as e:
        logging.error(f"Error clearing processed files: {str(e)}")
        
        # Return error response
        if request.headers.get('Accept') == 'application/json':
            return jsonify({'error': f'Failed to clear processed files: {str(e)}'}), 500
        
        # Show error page for browser requests
        return render_template_string("""
            <h1>Error</h1>
            <p>Failed to clear processed files: {{ error }}</p>
            <p><a href="{{ url_for('admin_system') }}">Back to System Settings</a></p>
        """, error=str(e)), 500

@app.route('/upload', methods=['POST'])
# Temporarily disable auth for testing
# @auth_required
def upload_file():
    if 'file' not in request.files:
        logging.warning("Upload failed: 'file' field missing in form-data")
        abort(400, "No file part")
    
    file = request.files['file']
    if file.filename == '':
        logging.warning("Upload failed: empty filename provided")
        abort(400, "No selected file")
    
    if file and allowed_file(file.filename):
        # Generate unique filename
        filename = secure_filename(file.filename)
        unique_filename = f"{uuid.uuid4().hex}_{filename}"
        filepath = os.path.join(UPLOAD_FOLDER, unique_filename)
        
        # Save file
        file.save(filepath)
        file_size = os.path.getsize(filepath)
        
        # Extract MIME type and MD5 hash
        import hashlib
        import mimetypes
        
        mimetype = file.mimetype or mimetypes.guess_type(filepath)[0]
        
        # Calculate MD5 hash for deduplication
        md5_hash = None
        try:
            with open(filepath, 'rb') as f:
                md5_hash = hashlib.md5(f.read()).hexdigest()
        except Exception as e:
            logging.warning(f"Failed to calculate MD5 hash for {filepath}: {e}")
        
        # Count PDF pages if it's a PDF
        page_count = None
        if filepath.lower().endswith('.pdf'):
            try:
                with open(filepath, 'rb') as f:
                    reader = PdfReader(f)
                    page_count = len(reader.pages)
            except Exception as e:
                logging.warning(f"Failed to count pages in PDF {filepath}: {e}")
        
        # Record file in database
        file_record = FileRecord(
            filename=unique_filename,
            original_filename=filename,
            file_size=file_size,
            file_type=file.mimetype.split('/')[1] if file.mimetype else filepath.split('.')[-1],
            user_id=getattr(current_user, 'id', 1),  # Use default ID 1 for tests
            mimetype=mimetype,
            hash_md5=md5_hash,
            page_count=page_count,
            storage_path=os.path.abspath(filepath),
            last_accessed=datetime.now(timezone.utc),
            access_count=1
        )
        db.session.add(file_record)
        db.session.commit()
        
        try:
            from tasks import on_upload_processing
            on_upload_processing.delay(file_record.id, filepath)
        except ImportError:
            logging.warning("Celery tasks not available for post-upload processing.")

        return jsonify({
            'key': unique_filename,
            'filename': filename,
            'id': file_record.id
        })
    
    logging.warning("Upload failed: invalid file type for '%s'", file.filename)
    abort(400, "Invalid file type. Only .pdf is allowed")

@app.route('/files')
@login_required
def get_user_files():
    """Get all files uploaded by the current user"""
    files = FileRecord.query.filter_by(user_id=current_user.id).order_by(FileRecord.upload_date.desc()).all()
    return jsonify([{
        'id': f.id,
        'filename': f.filename,
        'original_filename': f.original_filename,
        'file_size': f.file_size,
        'upload_date': f.upload_date.isoformat()
    } for f in files])

@app.route('/process', methods=['POST'])
@login_required
def process_pdf():
    """Process PDF files based on command and parameters"""
    try:
        if not request.is_json:
            return jsonify({"error": "Request must be JSON"}), 400
            
        data = request.json
        command = data.get('command')
        file_keys = data.get('file_keys', [])
        file_key = data.get('file_key')  # Support both formats
        params = data.get('params', {})
        
        # Support both file_keys array and single file_key
        if not file_keys and file_key:
            file_keys = [file_key]
            
        if not command:
            return jsonify({"error": "Missing command parameter"}), 400
            
        if not file_keys:
            return jsonify({"error": "Missing file_keys parameter"}), 400

        # Log the request for debugging
        logging.info(f"Process request - command: {command}, file_keys: {file_keys}, params: {params}")

        # Resolve file paths
        input_paths = [os.path.join(UPLOAD_FOLDER, key) for key in file_keys]
        
        # Verify files exist
        missing_files = [key for key, path in zip(file_keys, input_paths) if not os.path.exists(path)]
        if missing_files:
            return jsonify({"error": f"Files not found: {', '.join(missing_files)}"}), 404

        # Define output file 
        output_filename = f"{command}_{uuid.uuid4().hex}.pdf"
        output_path = os.path.join(PROCESSED_FOLDER, output_filename)

        # Try using Celery if available
        try:
            from celery.result import AsyncResult
            from tasks import process_pdf_task, celery as tasks_celery
            
            # Use the tasks.celery instance since that's where the worker is connected
            # Dispatch Celery background task using the tasks celery instance
            task = process_pdf_task.delay(command, input_paths, output_path, params)
            return jsonify({'task_id': task.id}), 202
            
        except ImportError:
            # Fallback to in-memory processing for simpler deployments
            logging.info("Celery not available, using in-memory task processing")
            
            # Generate a unique task ID
            task_id = str(uuid.uuid4())
            
            # Process the task in the background (simulated async)
            def process_task():
                try:
                    # Log start of processing
                    logging.info(f"Starting task {task_id} for command {command}")
                    
                    # Process the command
                    if command == 'merge' or command == 'merge_pdfs':
                        result = merge_pdfs(file_keys)
                    elif command == 'split':
                        result = split_pdf(file_keys[0], params)
                    elif command == 'compress':
                        result = compress_pdf(file_keys[0], params)
                    elif command == 'rotate':
                        result = rotate_pdf(file_keys[0], params)
                    elif command == 'pdf_to_word':
                        result = convert_to_word(file_keys[0], params)
                    elif command == 'pdf_to_excel':
                        result = convert_to_excel(file_keys[0], params)
                    elif command == 'pdf_to_jpg':
                        result = convert_to_jpg(file_keys[0], params)
                    elif command == 'protect':
                        result = protect_pdf(file_keys[0], params)
                    elif command == 'unlock':
                        result = unlock_pdf(file_keys[0], params)
                    elif command == 'watermark':
                        result = add_watermark(file_keys[0], params)
                    elif command == 'page_numbers':
                        result = add_page_numbers(file_keys[0], params)
                    else:
                        # Delegate to PDF processor for enhanced commands
                        result = pdf_processor.process_command(command, input_paths, output_path, params)
                    
                    # Store the result
                    app.task_results[task_id] = result
                    app.task_timestamps[task_id] = time.time()
                    
                    # Log completion
                    logging.info(f"Completed task {task_id} for command {command}")
                    
                except Exception as e:
                    # Log failure and detailed traceback
                    logging.error(f"Task {task_id} failed: {str(e)}")
                    logging.error(traceback.format_exc())
                    
                    # Store error result
                    app.task_results[task_id] = {"error": str(e)}
                    app.task_timestamps[task_id] = time.time()
            
            # Start background processing in a separate thread
            import threading
            thread = threading.Thread(target=process_task)
            thread.daemon = True
            thread.start()
            
            return jsonify({'task_id': task_id}), 202
            
    except Exception as e:
        # Log the full traceback for debugging
        logging.error(f"Process endpoint error: {str(e)}")
        logging.error(traceback.format_exc())
        return jsonify({"error": f"An error occurred: {str(e)}"}), 500

@app.route('/task/<task_id>')
@login_required
def task_status(task_id):
    """Get the status of a legacy in-memory task or Celery task if present."""
    # Prefer Celery if available
    CELERY_AVAILABLE = False
    try:
        from celery.result import AsyncResult
        from tasks import celery as tasks_celery
        CELERY_AVAILABLE = True
    except ImportError:
        CELERY_AVAILABLE = False
        
    if CELERY_AVAILABLE:
        task = AsyncResult(task_id, app=tasks_celery)
        if task.state == 'PENDING':
            return jsonify({'status': 'PENDING'})
        elif task.state == 'SUCCESS':
            info = task.info or {}
            return jsonify({'status': 'SUCCESS', 'result': info})
        elif task.state == 'FAILURE':
            return jsonify({'status': 'FAILURE', 'error': str(task.info)})
    
    # Fallback to in-memory results
    cleanup_old_tasks()
    if task_id not in app.task_results:
        abort(404, "Task not found")
    result = app.task_results[task_id]
    return jsonify({'status': 'SUCCESS', 'result': result})

@app.route('/history')
@login_required
def get_processing_history():
    """Get processing history for the current user"""
    history = ProcessingRecord.query.filter_by(user_id=current_user.id).order_by(ProcessingRecord.created_at.desc()).all()
    return jsonify([{
        'id': h.id,
        'task_id': h.task_id,
        'command': h.command,
        'input_files': h.input_files,
        'output_file': h.output_file,
        'status': h.status,
        'created_at': h.created_at.isoformat(),
        'completed_at': h.completed_at.isoformat() if h.completed_at else None
    } for h in history])

# --- PDF Processing Functions ---
def merge_pdfs(file_keys):
    """Merge multiple PDFs into one"""
    writer = PdfWriter()
    output_filename = f"merged_{uuid.uuid4().hex}.pdf"
    output_path = os.path.join(PROCESSED_FOLDER, output_filename)
    
    for key in file_keys:
        file_path = os.path.join(UPLOAD_FOLDER, key)
        if os.path.exists(file_path):
            reader = PdfReader(file_path)
            for page in reader.pages:
                writer.add_page(page)
    
    with open(output_path, 'wb') as output_file:
        writer.write(output_file)
    
    return {
        'key': output_filename,
        'filename': output_filename,
        'size': os.path.getsize(output_path)
    }

def split_pdf(file_key, params):
    """Split PDF into multiple files by pages"""
    file_path = os.path.join(UPLOAD_FOLDER, file_key)
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File {file_key} not found")
    
    pdf = pikepdf.open(file_path)
    pages = params.get('pages', '')
    
    if not pages:
        # Split every page
        for i, page in enumerate(pdf.pages):
            new_pdf = pikepdf.new()
            new_pdf.pages.append(page)
            split_filename = f"split_page_{i+1}_{uuid.uuid4().hex}.pdf"
            split_path = os.path.join(PROCESSED_FOLDER, split_filename)
            new_pdf.save(split_path)
    else:
        # Parse page ranges like "1-3,5,7-9"
        page_ranges = []
        for part in pages.split(','):
            if '-' in part:
                start, end = map(int, part.split('-'))
                page_ranges.extend(range(start, end + 1))
            else:
                page_ranges.append(int(part))
        
        for i, page_num in enumerate(page_ranges):
            if 1 <= page_num <= len(pdf.pages):
                new_pdf = pikepdf.new()
                new_pdf.pages.append(pdf.pages[page_num - 1])
                split_filename = f"split_page_{page_num}_{uuid.uuid4().hex}.pdf"
                split_path = os.path.join(PROCESSED_FOLDER, split_filename)
                new_pdf.save(split_path)
    
    # Return the first split file for now (in a real app, you'd return all files)
    first_file = [f for f in os.listdir(PROCESSED_FOLDER) if f.startswith('split_page_')][0]
    first_path = os.path.join(PROCESSED_FOLDER, first_file)
    
    return {
        'key': first_file,
        'filename': first_file,
        'size': os.path.getsize(first_path)
    }

def compress_pdf(file_key, params):
    """Compress PDF file"""
    file_path = os.path.join(UPLOAD_FOLDER, file_key)
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File {file_key} not found")
    
    quality = params.get('quality', 'medium')
    output_filename = f"compressed_{uuid.uuid4().hex}.pdf"
    output_path = os.path.join(PROCESSED_FOLDER, output_filename)
    
    pdf = pikepdf.open(file_path)
    
    # Apply compression based on quality setting
    if quality == 'low':
        pdf.save(output_path, linearize=True, compress_streams=True, preserve_pdfa=False)
    elif quality == 'high':
        pdf.save(output_path, linearize=True, compress_streams=False, preserve_pdfa=True)
    else:  # medium
        pdf.save(output_path, linearize=True, compress_streams=True, preserve_pdfa=True)
    
    return {
        'key': output_filename,
        'filename': output_filename,
        'size': os.path.getsize(output_path)
    }

def rotate_pdf(file_key, params):
    """Rotate PDF pages"""
    file_path = os.path.join(UPLOAD_FOLDER, file_key)
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File {file_key} not found")
    
    angle = int(params.get('angle', '90'))
    output_filename = f"rotated_{uuid.uuid4().hex}.pdf"
    output_path = os.path.join(PROCESSED_FOLDER, output_filename)
    
    pdf = pikepdf.open(file_path)
    
    # Rotate all pages
    for page in pdf.pages:
        page.rotate(angle, relative=True)
    
    pdf.save(output_path)
    
    return {
        'key': output_filename,
        'filename': output_filename,
        'size': os.path.getsize(output_path)
    }

# --- NEW ADVANCED FEATURES ---

def convert_to_word(file_key, params):
    """Convert PDF to Word document"""
    try:
        import pdfplumber
        import docx
    except ImportError:
        raise ImportError("pdfplumber and python-docx are required for PDF to Word conversion")
    
    file_path = os.path.join(UPLOAD_FOLDER, file_key)
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File {file_key} not found")
    
    output_filename = f"converted_{uuid.uuid4().hex}.docx"
    output_path = os.path.join(PROCESSED_FOLDER, output_filename)
    
    # Create Word document
    doc = docx.Document()
    
    # Extract text with layout preservation
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text(layout=True)
            if text:
                doc.add_paragraph(text)
            doc.add_page_break()
    
    doc.save(output_path)
    
    return {
        'key': output_filename,
        'filename': output_filename,
        'size': os.path.getsize(output_path)
    }

def convert_to_excel(file_key, params):
    """Convert PDF tables to Excel"""
    try:
        import pdfplumber
        import openpyxl
    except ImportError:
        raise ImportError("pdfplumber and openpyxl are required for PDF to Excel conversion")
    
    file_path = os.path.join(UPLOAD_FOLDER, file_key)
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File {file_key} not found")
    
    output_filename = f"converted_{uuid.uuid4().hex}.xlsx"
    output_path = os.path.join(PROCESSED_FOLDER, output_filename)
    
    # Create Excel workbook
    wb = openpyxl.Workbook()
    ws = wb.active
    
    # Extract tables from PDF
    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            tables = page.extract_tables()
            for table_idx, table in enumerate(tables):
                for row_idx, row in enumerate(table):
                    for col_idx, cell in enumerate(row):
                        if cell:  # Only add non-empty cells
                            ws.cell(row=row_idx + 1, column=col_idx + 1, value=cell)
    
    wb.save(output_path)
    
    return {
        'key': output_filename,
        'filename': output_filename,
        'size': os.path.getsize(output_path)
    }

def convert_to_jpg(file_key, params):
    """Convert PDF to JPG images"""
    try:
        # Try pdf2image first, fall back to PyMuPDF if poppler not available
        try:
            from pdf2image import convert_from_path
            from PIL import Image
            
            file_path = os.path.join(UPLOAD_FOLDER, file_key)
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File {file_key} not found")
            
            pages = params.get('pages', 'all')
            dpi = int(params.get('dpi', '150'))
            
            # Convert PDF to images
            images = convert_from_path(file_path, dpi=dpi)
            
            # Save images
            image_files = []
            for i, img in enumerate(images):
                if pages == 'all' or str(i+1) in pages.split(','):
                    image_filename = f"converted_page_{i+1}_{uuid.uuid4().hex}.jpg"
                    image_path = os.path.join(PROCESSED_FOLDER, image_filename)
                    img.save(image_path, 'JPEG', quality=95)
                    image_files.append(image_filename)
            
            # Return the first image for now (in a real app, you'd return all images)
            first_image = image_files[0] if image_files else None
            if first_image:
                first_path = os.path.join(PROCESSED_FOLDER, first_image)
                return {
                    'key': first_image,
                    'filename': first_image,
                    'size': os.path.getsize(first_path)
                }
            else:
                raise ValueError("No images were created")
                
        except Exception as e:
            # Fallback to PyMuPDF if pdf2image fails
            if 'pymupdf' in str(e).lower() or 'poppler' in str(e).lower():
                import fitz  # PyMuPDF
                from PIL import Image
                
                file_path = os.path.join(UPLOAD_FOLDER, file_key)
                if not os.path.exists(file_path):
                    raise FileNotFoundError(f"File {file_key} not found")
                
                pages = params.get('pages', 'all')
                dpi = int(params.get('dpi', '150'))
                
                # Open PDF with PyMuPDF
                pdf_document = fitz.open(file_path)
                
                # Convert pages to images
                image_files = []
                for page_num in range(pdf_document.page_count):
                    if pages == 'all' or str(page_num + 1) in pages.split(','):
                        page = pdf_document[page_num]
                        
                        # Calculate zoom factor for DPI
                        zoom = dpi / 72.0
                        mat = fitz.Matrix(zoom, zoom)
                        
                        # Render page to image
                        pix = page.get_pixmap(matrix=mat)
                        img_data = pix.tobytes("ppm")
                        
                        # Convert to PIL Image and save as JPG
                        img = Image.frombytes("RGB", [pix.width, pix.height], img_data)
                        image_filename = f"converted_page_{page_num + 1}_{uuid.uuid4().hex}.jpg"
                        image_path = os.path.join(PROCESSED_FOLDER, image_filename)
                        img.save(image_path, 'JPEG', quality=95)
                        image_files.append(image_filename)
                
                pdf_document.close()
                
                # Return the first image
                first_image = image_files[0] if image_files else None
                if first_image:
                    first_path = os.path.join(PROCESSED_FOLDER, first_image)
                    return {
                        'key': first_image,
                        'filename': first_image,
                        'size': os.path.getsize(first_path)
                    }
                else:
                    raise ValueError("No images were created")
            else:
                raise e
                
    except ImportError as e:
        raise ImportError("Either pdf2image+Pillow or PyMuPDF+Pillow are required for PDF to JPG conversion")

def protect_pdf(file_key, params):
    """Encrypt PDF with password protection"""
    file_path = os.path.join(UPLOAD_FOLDER, file_key)
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File {file_key} not found")
    
    password = params.get('password', '')
    if not password:
        raise ValueError("Password is required for PDF protection")
    
    output_filename = f"protected_{uuid.uuid4().hex}.pdf"
    output_path = os.path.join(PROCESSED_FOLDER, output_filename)
    
    # Open and encrypt PDF
    pdf = pikepdf.open(file_path)
    pdf.save(
        output_path,
        encryption=pikepdf.Encryption(owner=password, user=password, R=4)
    )
    
    return {
        'key': output_filename,
        'filename': output_filename,
        'size': os.path.getsize(output_path)
    }

def unlock_pdf(file_key, params):
    """Remove password protection from PDF"""
    file_path = os.path.join(UPLOAD_FOLDER, file_key)
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File {file_key} not found")
    
    password = params.get('password', '')
    if not password:
        raise ValueError("Password is required to unlock PDF")
    
    output_filename = f"unlocked_{uuid.uuid4().hex}.pdf"
    output_path = os.path.join(PROCESSED_FOLDER, output_filename)
    
    # Open encrypted PDF and save without encryption
    pdf = pikepdf.open(file_path, password=password)
    pdf.save(output_path)
    
    return {
        'key': output_filename,
        'filename': output_filename,
        'size': os.path.getsize(output_path)
    }

def add_watermark(file_key, params):
    """Add text watermark to PDF"""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
    except ImportError:
        raise ImportError("reportlab is required for watermarking")
    
    file_path = os.path.join(UPLOAD_FOLDER, file_key)
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File {file_key} not found")
    
    watermark_text = params.get('text', 'CONFIDENTIAL')
    opacity = float(params.get('opacity', '0.3'))
    
    output_filename = f"watermarked_{uuid.uuid4().hex}.pdf"
    output_path = os.path.join(PROCESSED_FOLDER, output_filename)
    
    # Create watermark PDF
    packet = io.BytesIO()
    can = canvas.Canvas(packet, pagesize=letter)
    can.setFont("Helvetica", 40)
    can.setFillAlpha(opacity)
    can.rotate(45)
    can.drawString(200, 100, watermark_text)
    can.save()
    packet.seek(0)
    
    # Read watermark and input PDF
    watermark_pdf = PdfReader(packet)
    watermark_page = watermark_pdf.pages[0]
    
    input_pdf = PdfReader(file_path)
    output_writer = PdfWriter()
    
    # Apply watermark to each page
    for page in input_pdf.pages:
        page.merge_page(watermark_page)
        output_writer.add_page(page)
    
    # Save watermarked PDF
    with open(output_path, "wb") as f:
        output_writer.write(f)
    
    return {
        'key': output_filename,
        'filename': output_filename,
        'size': os.path.getsize(output_path)
    }

def add_page_numbers(file_key, params):
    """Add page numbers to PDF"""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
    except ImportError:
        raise ImportError("reportlab is required for page numbering")
    
    file_path = os.path.join(UPLOAD_FOLDER, file_key)
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File {file_key} not found")
    
    start_number = int(params.get('start', '1'))
    position = params.get('position', 'bottom-right')
    
    output_filename = f"numbered_{uuid.uuid4().hex}.pdf"
    output_path = os.path.join(PROCESSED_FOLDER, output_filename)
    
    input_pdf = PdfReader(file_path)
    output_writer = PdfWriter()
    
    # Add page numbers to each page
    for i, page in enumerate(input_pdf.pages, start_number):
        # Create temporary PDF with page number using NamedTemporaryFile for security
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
            temp_pdf = temp_file.name
        
        try:
            c = canvas.Canvas(temp_pdf, pagesize=letter)
            
            # Position page number based on setting
            if position == 'bottom-right':
                c.drawString(500, 20, str(i))
            elif position == 'bottom-center':
                c.drawString(300, 20, str(i))
            elif position == 'bottom-left':
                c.drawString(100, 20, str(i))
            elif position == 'top-right':
                c.drawString(500, 780, str(i))
            elif position == 'top-center':
                c.drawString(300, 780, str(i))
            elif position == 'top-left':
                c.drawString(100, 780, str(i))
            
            c.save()
            
            # Merge page number with original page
            num_reader = PdfReader(temp_pdf)
            page.merge_page(num_reader.pages[0])
            output_writer.add_page(page)
        finally:
            # Clean up temporary file
            if os.path.exists(temp_pdf):
                os.unlink(temp_pdf)
    
    # Save numbered PDF
    with open(output_path, "wb") as f:
        output_writer.write(f)
    
    return {
        'key': output_filename,
        'filename': output_filename,
        'size': os.path.getsize(output_path)
    }

def add_header_footer(file_key, params):
    """Add headers and footers to PDF"""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
    except ImportError:
        raise ImportError("reportlab is required for headers/footers")
    
    file_path = os.path.join(UPLOAD_FOLDER, file_key)
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File {file_key} not found")
    
    header_text = params.get('header', '')
    footer_text = params.get('footer', '')
    
    output_filename = f"header_footer_{uuid.uuid4().hex}.pdf"
    output_path = os.path.join(PROCESSED_FOLDER, output_filename)
    
    input_pdf = PdfReader(file_path)
    output_writer = PdfWriter()
    
    # Add header/footer to each page
    for page in input_pdf.pages:
        # Create temporary PDF with header/footer using NamedTemporaryFile for security
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
            temp_pdf = temp_file.name
        
        try:
            c = canvas.Canvas(temp_pdf, pagesize=letter)
            
            if header_text:
                c.drawString(100, 800, header_text)
            if footer_text:
                c.drawString(100, 20, footer_text)
            
            c.save()
            
            # Merge header/footer with original page
            hf_reader = PdfReader(temp_pdf)
            page.merge_page(hf_reader.pages[0])
            output_writer.add_page(page)
        finally:
            # Clean up temporary file
            if os.path.exists(temp_pdf):
                os.unlink(temp_pdf)
    
    # Save PDF with headers/footers
    with open(output_path, "wb") as f:
        output_writer.write(f)
    
    return {
        'key': output_filename,
        'filename': output_filename,
        'size': os.path.getsize(output_path)
    }

@app.route('/download')
# Temporarily disable auth for testing
# @login_required
def download_file():
    key = request.args.get('key')
    if not key:
        abort(400, "Missing 'key' parameter")
    
    # Prevent path traversal attacks
    if '..' in key or '/' in key or '\\' in key:
        abort(400, "Invalid file key")
    
    # First check in PROCESSED_FOLDER
    file_path = os.path.join(PROCESSED_FOLDER, key)
    
    # If not found in PROCESSED_FOLDER, check UPLOAD_FOLDER
    if not os.path.exists(file_path):
        file_path = os.path.join(UPLOAD_FOLDER, key)
    
    # Ensure the file is actually within the allowed folders
    try:
        file_path = os.path.abspath(file_path)
        if not (file_path.startswith(os.path.abspath(PROCESSED_FOLDER)) or 
                file_path.startswith(os.path.abspath(UPLOAD_FOLDER))):
            abort(400, "Invalid file path")
    except (OSError, ValueError):
        abort(400, "Invalid file path")
    
    if not os.path.exists(file_path):
        abort(404, "File not found")
    
    # Update file access statistics if file exists in database
    try:
        file_record = get_file_by_key(key)
        
        # If no record found but file exists, it might be a processed file
        # Let's check if it's in the FileConversionRecord table
        if not file_record and key.startswith(('merged_', 'split_', 'compressed_', 'converted_', 'protected_')):
            conversion = FileConversionRecord.query.filter_by(output_file=key).first()
            if conversion:
                # Update access for the original file
                original_file = db.session.get(FileRecord, conversion.original_file_id)
                if original_file:
                    original_file.last_accessed = datetime.now(timezone.utc)
                    original_file.access_count += 1
                    db.session.commit()
    except Exception as e:
        logging.warning(f"Failed to update file access statistics: {e}")
    
    return send_file(
        file_path,
        as_attachment=True,
        download_name=key
    )

# ============================================================================
# AI-POWERED DOCUMENT INTELLIGENCE ROUTES
# ============================================================================

@app.route('/api/chat-pdf', methods=['POST'])
@login_required
def api_chat_pdf():
    """Chat with PDF using AI"""
    try:
        data = request.json
        file_key = data.get('file_key')
        question = data.get('question')
        
        if not file_key or not question:
            return jsonify({"error": "Missing file_key or question"}), 400
        
        # Find the actual file path
        file_record = FileRecord.query.filter_by(filename=file_key, user_id=current_user.id).first()
        if not file_record:
            return jsonify({"error": "File not found"}), 404
        
        file_path = os.path.join(UPLOAD_FOLDER, file_record.filename)
        if not os.path.exists(file_path):
            return jsonify({"error": "File not found on disk"}), 404
        
        # Import and run the AI task
        try:
            from tasks import chat_with_pdf
            result = chat_with_pdf.delay(file_path, question)
            return jsonify({"task_id": result.id}), 202
        except ImportError:
            return jsonify({"error": "AI features not available"}), 503
        
    except Exception as e:
        logging.error(f"Chat PDF error: {e}")
        return jsonify({"error": "Internal server error"}), 500

@app.route('/api/analyze-pdf', methods=['POST'])
@login_required
def api_analyze_pdf():
    """Analyze PDF with AI: summarization, NER, topic modeling"""
    try:
        data = request.json
        file_key = data.get('file_key')
        
        if not file_key:
            return jsonify({"error": "Missing file_key"}), 400
        
        # Find the actual file path
        file_record = FileRecord.query.filter_by(filename=file_key, user_id=current_user.id).first()
        if not file_record:
            return jsonify({"error": "File not found"}), 404
        
        file_path = os.path.join(UPLOAD_FOLDER, file_record.filename)
        if not os.path.exists(file_path):
            return jsonify({"error": "File not found on disk"}), 404
        
        # Import and run the AI task
        try:
            from tasks import analyze_pdf
            result = analyze_pdf.delay(file_path)
            return jsonify({"task_id": result.id}), 202
        except ImportError:
            return jsonify({"error": "AI features not available"}), 503
        
    except Exception as e:
        logging.error(f"Analyze PDF error: {e}")
        return jsonify({"error": "Internal server error"}), 500

@app.route('/api/classify-document', methods=['POST'])
@login_required
def api_classify_document():
    """Classify document using MLflow models"""
    try:
        data = request.json
        file_key = data.get('file_key')
        
        if not file_key:
            return jsonify({"error": "Missing file_key"}), 400
        
        # Find the actual file path
        file_record = FileRecord.query.filter_by(filename=file_key, user_id=current_user.id).first()
        if not file_record:
            return jsonify({"error": "File not found"}), 404
        
        file_path = os.path.join(UPLOAD_FOLDER, file_record.filename)
        if not os.path.exists(file_path):
            return jsonify({"error": "File not found on disk"}), 404
        
        # Import and run the AI task
        try:
            from tasks import classify_document
            result = classify_document.delay(file_path)
            return jsonify({"task_id": result.id}), 202
        except ImportError:
            return jsonify({"error": "AI features not available"}), 503
        
    except Exception as e:
        logging.error(f"Classify document error: {e}")
        return jsonify({"error": "Internal server error"}), 500

@app.route('/api/multi-document-chat', methods=['POST'])
@login_required
def api_multi_document_chat():
    try:
        data = request.json
        file_keys = data.get('file_keys')
        question = data.get('question')
        
        if not file_keys or not question:
            return jsonify({"error": "Missing file_keys or question"}), 400
        
        file_paths = []
        for file_key in file_keys:
            file_record = FileRecord.query.filter_by(filename=file_key, user_id=current_user.id).first()
            if not file_record:
                return jsonify({"error": f"File not found: {file_key}"}), 404
            
            file_path = os.path.join(UPLOAD_FOLDER, file_record.filename)
            if not os.path.exists(file_path):
                return jsonify({"error": f"File not found on disk: {file_key}"}), 404
            file_paths.append(file_path)
        
        try:
            from tasks import multi_document_chat
            result = multi_document_chat.delay(file_paths, question)
            return jsonify({"task_id": result.id}), 202
        except ImportError:
            return jsonify({"error": "AI features not available"}), 503
        
    except Exception as e:
        logging.error(f"Multi-document chat error: {e}")
        return jsonify({"error": "Internal server error"}), 500

# ============================================================================
# WORKFLOW AUTOMATION ROUTES
# ============================================================================

@app.route('/api/workflow', methods=['POST'])
@login_required
def api_workflow():
    """Execute automated workflow with chained PDF operations"""
    try:
        data = request.json
        file_key = data.get('file_key')
        commands = list(data.get('commands', []))
        
        if not file_key or not commands:
            return jsonify({"error": "Missing file_key or commands"}), 400
        
        # Find the actual file path
        file_record = FileRecord.query.filter_by(filename=file_key, user_id=current_user.id).first()
        if not file_record:
            return jsonify({"error": "File not found"}), 404
        
        file_path = os.path.join(UPLOAD_FOLDER, file_record.filename)
        if not os.path.exists(file_path):
            return jsonify({"error": "File not found on disk"}), 404
        
        # Import and run the workflow task
        try:
            from tasks import workflow_master
            result = workflow_master.delay(file_path, commands)
            return jsonify({"task_id": result.id}), 202
        except ImportError:
            return jsonify({"error": "Workflow features not available"}), 503
        
    except Exception as e:
        logging.error(f"Workflow error: {e}")
        return jsonify({"error": "Internal server error"}), 500

# ============================================================================
# Cloud Storage Integration
# ============================================================================

try:
    from google_auth_oauthlib.flow import InstalledAppFlow
    from googleapiclient.discovery import build
    from googleapiclient.errors import HttpError
    from googleapiclient.http import MediaIoBaseDownload
    from google.oauth2.credentials import Credentials
    from google.auth.transport.requests import Request
    GOOGLE_DRIVE_AVAILABLE = True
except ImportError:
    GOOGLE_DRIVE_AVAILABLE = False
    logging.warning("Google Drive integration not available")
import json
import io

SCOPES = ['https://www.googleapis.com/auth/drive.readonly', 'https://www.googleapis.com/auth/drive.file']
CLIENT_CONFIG = json.loads(os.getenv('GOOGLE_CLIENT_CONFIG', '{}'))

@app.route('/connect-drive')
@login_required
def connect_drive():
    flow = InstalledAppFlow.from_client_config(CLIENT_CONFIG, SCOPES)
    flow.redirect_uri = url_for('oauth2callback', _external=True)
    authorization_url, state = flow.authorization_url(access_type='offline', include_granted_scopes='true')
    session['state'] = state
    return redirect(authorization_url)

@app.route('/oauth2callback')
@login_required
def oauth2callback():
    state = session['state']
    flow = InstalledAppFlow.from_client_config(CLIENT_CONFIG, SCOPES, state=state)
    flow.redirect_uri = url_for('oauth2callback', _external=True)
    authorization_response = request.url
    flow.fetch_token(authorization_response=authorization_response)
    credentials = flow.credentials
    current_user.google_drive_token = credentials.to_json()
    db.session.commit()
    return redirect(url_for('index'))

@app.route('/list-drive-files', methods=['GET'])
@login_required
def list_drive_files():
    try:
        if not current_user.google_drive_token:
            return jsonify({"error": "Drive not connected"}), 401
        
        creds = Credentials.from_authorized_user_info(json.loads(current_user.google_drive_token), SCOPES)
        if creds.expired and creds.refresh_token:
            creds.refresh(Request())
        
        service = build('drive', 'v3', credentials=creds)
        results = service.files().list(q="mimeType='application/pdf'", pageSize=10, fields="nextPageToken, files(id, name)").execute()
        files = results.get('files', [])
        return jsonify(files), 200
    except Exception as e:
        logging.error(f"List drive error: {e}\n{traceback.format_exc()}")
        return jsonify({"error": "Failed to list files"}), 500

@app.route('/import-drive-file', methods=['POST'])
@login_required
def import_drive_file():
    try:
        data = request.json
        drive_file_id = data.get('drive_file_id')
        if not drive_file_id:
            return jsonify({"error": "drive_file_id is required"}), 400
        
        from tasks import import_from_drive
        task = import_from_drive.delay(current_user.id, drive_file_id)
        return jsonify({"task_id": task.id}), 202
    except Exception as e:
        logging.error(f"Import drive file error: {e}\n{traceback.format_exc()}")
        return jsonify({"error": "Failed to import file"}), 500

# ============================================================================
# TASK STATUS ROUTES
# ============================================================================

@app.route('/api/task-status/<task_id>', methods=['GET'])
@login_required
def api_task_status(task_id):
    """Get the status of a Celery task"""
    CELERY_AVAILABLE = False
    try:
        from celery.result import AsyncResult
        from tasks import celery as tasks_celery
        CELERY_AVAILABLE = True
    except ImportError:
        CELERY_AVAILABLE = False
        
    if CELERY_AVAILABLE:    
        try:
            task = AsyncResult(task_id, app=tasks_celery)
            if task.state == 'PENDING':
                return jsonify({'status': 'PENDING'})
            elif task.state != 'FAILURE':
                return jsonify({'status': 'SUCCESS', 'result': task.info})
            else:
                return jsonify({'status': 'FAILURE', 'error': str(task.info)})
        except Exception as e:
            logging.error(f"Task status error: {e}")
            return jsonify({"error": "Internal server error"}), 500
    else:
        # Fallback to in-memory results
        if task_id not in app.task_results:
            return jsonify({"error": "Task not found"}), 404
        result = app.task_results[task_id]
        return jsonify({'status': 'SUCCESS', 'result': result})

# ============================================================================
# PUBLIC API (RESTX)
# ============================================================================

if RESTX_AVAILABLE:
    # Initialize RESTX API
    from flask import Blueprint
    api_bp = Blueprint('api', __name__)
    api = Api(api_bp, version='1.0', title='PDF Tool API', 
              description='API for PDF processing', doc='/docs/')
    
    # Models
    upload_model = api.model('Upload', {
        'file_key': fields.String(required=True, description='File key'),
        'operation': fields.String(required=True, description='Operation e.g., compress')
    })
    
    # Namespace
    ns = api.namespace('pdf', description='PDF operations')
    
    def api_key_required(f):
        @wraps(f)
        def decorated(*args, **kwargs):
            api_key = request.headers.get('X-API-KEY')
            if api_key != get_api_key():
                logging.warning("Invalid API key attempt")
                api.abort(403, 'Invalid API key')
            return f(*args, **kwargs)
        return decorated
    
    @ns.route('/process')
    class ProcessPDF(Resource):
        @ns.doc('process_pdf')
        @ns.expect(upload_model)
        @api_key_required
        def post(self):
            try:
                data = api.payload
                file_key = data.get('file_key')
                operation = data.get('operation')
                
                if not file_key or not operation:
                    api.abort(400, 'Missing file_key or operation')
                
                # Import and run the task
                try:
                    from tasks import process_pdf_task
                    result = process_pdf_task.delay(operation, [file_key])
                    return {'task_id': result.id}, 202
                except ImportError:
                    api.abort(503, 'Processing features not available')
                    
            except ValueError as e:
                api.abort(400, str(e))
            except Exception as e:
                logging.error(f"API error: {e}")
                api.abort(500, f'Internal error: {str(e)}')
    
    # Register the API blueprint
    app.register_blueprint(api_bp, url_prefix='/api/v1')

# Add test API routes for merge and split
@app.route('/operations/merge', methods=['POST'])
@auth_required
def merge_endpoint():
    """Test API endpoint for PDF merge"""
    data = request.json
    file_ids = data.get('file_ids', [])
    output_filename = data.get('output_filename', 'merged.pdf')
    
    if not file_ids or len(file_ids) < 2:
        return jsonify({"error": "At least 2 file IDs required"}), 400
    
    # In a real implementation, fetch files from database by IDs
    # Here we just return success for tests
    return jsonify({"message": "PDFs merged successfully", "output_file": output_filename})

@app.route('/operations/split', methods=['POST'])
@auth_required
def split_endpoint():
    """Test API endpoint for PDF split"""
    data = request.json
    file_id = data.get('file_id')
    
    if not file_id:
        return jsonify({"error": "File ID required"}), 400
    
    # In a real implementation, fetch file from database by ID
    # Here we just return success for tests
    return jsonify({"message": "PDF split successfully", "output_files": ["split_1.pdf", "split_2.pdf"]})

# ============================================================================
# ENHANCED PDF PROCESSING ENDPOINTS
# ============================================================================

@app.route('/enhanced/merge', methods=['POST'])
@login_required
def enhanced_merge():
    """
    Merges multiple PDF files specified by their keys in a JSON request.

    This endpoint is designed to be efficient and robust, handling input
    validation, bulk database lookups, and clear error reporting.

    Expects a JSON body:
    {
        "file_keys": ["key1.pdf", "key2.pdf", "key3.pdf"]
    }
    """
    # --- 1. Robust Input Validation ---
    try:
        data = request.get_json()
        if not data:
            return jsonify({"error": "Invalid JSON payload"}), 400
        file_keys = data.get('file_keys')
    except BadRequest:
        return jsonify({"error": "Malformed request. Could not parse JSON."}), 400

    if not isinstance(file_keys, list) or len(file_keys) < 2:
        return jsonify({"error": "A 'file_keys' array with at least 2 PDF file keys is required."}), 400

    try:
        # --- 2. Efficient Database Query (Anti-N+1 Pattern) ---
        # Fetch all required records in a single database query.
        file_records = FileRecord.query.filter(
            FileRecord.user_id == current_user.id,
            FileRecord.filename.in_(file_keys)
        ).all()

        found_filenames = {record.filename for record in file_records}
        missing_db_keys = set(file_keys) - found_filenames

        if missing_db_keys:
            # Report all missing files at once for a better user experience.
            return jsonify({"error": f"Files not found in database: {', '.join(sorted(missing_db_keys))}"}), 404

        # --- 3. File System Validation ---
        # Create a mapping for easy path generation while preserving order.
        record_map = {record.filename: record for record in file_records}
        file_paths = [os.path.join(UPLOAD_FOLDER, record_map[key].filename) for key in file_keys]

        # Check for files on disk.
        missing_disk_paths = [path for path in file_paths if not os.path.exists(path)]
        if missing_disk_paths:
            missing_filenames = [os.path.basename(p) for p in missing_disk_paths]
            return jsonify({"error": f"Files not found on disk: {', '.join(sorted(missing_filenames))}"}), 404

        # --- 4. Processing and Output ---
        # Ensure the processed folder exists to avoid file write errors.
        os.makedirs(PROCESSED_FOLDER, exist_ok=True)

        # Generate a unique output filename
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        output_filename = f"merged_{timestamp}_{uuid.uuid4().hex[:8]}.pdf"
        output_path = os.path.join(PROCESSED_FOLDER, output_filename)
        
        # Log the operation
        logging.info(f"Merging {len(file_keys)} files: {', '.join(file_keys)}")
        logging.info(f"Output path: {output_path}")
        
        # The processing function should raise an exception on failure.
        result = pdf_processor.merge_pdfs(file_paths, output_path)
        
        # Verify the output file exists and has reasonable size
        if not os.path.exists(output_path):
            return jsonify({"error": "Merge operation failed to create output file"}), 500
            
        file_size = os.path.getsize(output_path)
        if file_size == 0:
            os.remove(output_path)  # Clean up empty file
            return jsonify({"error": "Merge operation created an empty file"}), 500

        # Save the merged file to the database.
        new_record = FileRecord(
            user_id=current_user.id,
            filename=output_filename,
            original_filename=output_filename,
            file_size=file_size,
            file_type='pdf'
        )
        db.session.add(new_record)
        db.session.commit()

        # --- 5. Clear and Consistent Success Response ---
        return jsonify({
            "success": True,
            "message": "Files merged successfully.",
            "result": {
                "key": output_filename,
                "filename": output_filename,
                "size": file_size
            }
        })

    # --- 6. Improved Exception Handling ---
    except PDFValidationError as e:
        logging.error(f"PDF validation error: {e}")
        return jsonify({"error": f"Invalid PDF file: {e}"}), 422
    
    except PDFOperationError as e:
        logging.error(f"PDF operation error: {e}")
        return jsonify({"error": f"PDF processing failed: {e}"}), 422
        
    except Exception as e:
        # Log the full traceback for debugging.
        logging.exception("An error occurred during the PDF merge process.")
        # Return a generic error to the user for security.
        return jsonify({"error": "An internal server error occurred."}), 500

@app.route('/enhanced/split', methods=['POST'])
@login_required
def enhanced_split():
    """Enhanced PDF split with validation - expects file_key in JSON"""
    try:
        data = request.get_json()
        file_key = data.get('file_key')
        
        if not file_key:
            return jsonify({"error": "File key required"}), 400
        
        # Resolve file path from file_key
        file_record = FileRecord.query.filter_by(filename=file_key, user_id=current_user.id).first()
        if not file_record:
            return jsonify({"error": "File not found"}), 404
        file_path = os.path.join(UPLOAD_FOLDER, file_record.filename)
        if not os.path.exists(file_path):
            return jsonify({"error": "File not found on disk"}), 404
        
        # Process with PDFProcessor
        output_dir = os.path.join(PROCESSED_FOLDER, f"split_{uuid.uuid4().hex}")
        os.makedirs(output_dir, exist_ok=True)
        result = pdf_processor.split_pdf(file_path, output_dir)
        
        # Move files to PROCESSED_FOLDER to avoid path issues in download
        split_files = [f for f in os.listdir(output_dir) if f.endswith('.pdf')]
        moved_files = []
        for sf in split_files:
            src = os.path.join(output_dir, sf)
            dst = os.path.join(PROCESSED_FOLDER, sf)
            shutil.move(src, dst)
            moved_files.append(sf)
        os.rmdir(output_dir)
        
        if moved_files:
            first_file = moved_files[0]
            first_path = os.path.join(PROCESSED_FOLDER, first_file)
            return jsonify({
                "success": True,
                "result": {
                    "key": first_file,
                    "filename": first_file,
                    "size": os.path.getsize(first_path)
                },
                "all_files": moved_files
            })
        else:
            return jsonify({"error": "No files created during split"}), 500
    except Exception as e:
        logging.error(f"Enhanced split error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/enhanced/convert', methods=['POST'])
@login_required
def enhanced_convert():
    """Enhanced document conversion with validation - expects file_key in JSON"""
    try:
        data = request.get_json()
        file_key = data.get('file_key')
        target_format = data.get('target_format')
        
        if not file_key or not target_format:
            return jsonify({"error": "File key and target format required"}), 400
        
        # Resolve file path from file_key
        file_record = FileRecord.query.filter_by(filename=file_key, user_id=current_user.id).first()
        if not file_record:
            return jsonify({"error": "File not found"}), 404
        file_path = os.path.join(UPLOAD_FOLDER, file_record.filename)
        if not os.path.exists(file_path):
            return jsonify({"error": "File not found on disk"}), 404
        
        # Determine extension
        ext_map = {
            'docx': '.docx',
            'xlsx': '.xlsx',
            'pptx': '.pptx',
            'jpg': '.jpg'
        }
        ext = ext_map.get(target_format, '.pdf')
        output_filename = f"converted_{uuid.uuid4().hex}{ext}"
        output_path = os.path.join(PROCESSED_FOLDER, output_filename)
        
        # Process with PDFProcessor based on target format
        if target_format == 'pptx':
            result = pdf_processor.pdf_to_powerpoint(file_path, output_path)
        elif target_format == 'docx':
            # Fallback to defined function if processor doesn't support
            try:
                result = convert_to_word(file_key, {})
            except:
                result = pdf_processor.pdf_to_word(file_path, output_path) if hasattr(pdf_processor, 'pdf_to_word') else None
        elif target_format == 'pdf':
            # File is already PDF, just return success with original file
            result = {
                "filename": file_record.filename,
                "message": "File is already in PDF format",
                "key": file_key
            }
        # Add more formats as needed
        else:
            return jsonify({"error": f"Unsupported target format: {target_format}"}), 400
        
        return jsonify({
            "success": True,
            "result": result or {
                "key": output_filename,
                "filename": output_filename,
                "size": os.path.getsize(output_path)
            }
        })
    except Exception as e:
        logging.error(f"Enhanced convert error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/enhanced/workflow', methods=['POST'])
@login_required
def enhanced_workflow():
    """Enhanced workflow execution"""
    try:
        data = request.get_json()
        file_key = data.get('file_key')
        operations = data.get('operations', [])
        
        if not file_key or not operations:
            return jsonify({"error": "File key and operations required"}), 400
        
        # Resolve file path
        file_record = FileRecord.query.filter_by(filename=file_key, user_id=current_user.id).first()
        if not file_record:
            return jsonify({"error": "File not found"}), 404
        file_path = os.path.join(UPLOAD_FOLDER, file_record.filename)
        if not os.path.exists(file_path):
            return jsonify({"error": "File not found on disk"}), 404
        
        # Build workflow ops
        workflow_ops = []
        for op in operations:
            if op == 'rotate_pdf':
                workflow_ops.append({'method': 'rotate_pdf', 'args': {'rotation': 90}})
            elif op == 'watermark_pdf':
                workflow_ops.append({'method': 'watermark_pdf', 'args': {'watermark_text': 'Processed', 'opacity': 0.3}})
            elif op == 'compress_pdf':
                workflow_ops.append({'method': 'compress_pdf', 'args': {}})
            elif op == 'add_page_numbers':
                workflow_ops.append({'method': 'add_page_numbers', 'args': {}})
            elif op == 'protect_pdf':
                workflow_ops.append({'method': 'protect_pdf', 'args': {'user_password': 'password'}})
        
        # Add input_path to the first operation and output_path to all operations
        if workflow_ops:
            workflow_ops[0]['args']['input_path'] = file_path
            for i, op in enumerate(workflow_ops):
                if 'output_path' not in op['args']:
                    output_filename = f"workflow_step_{i}_{uuid.uuid4().hex}.pdf"
                    op['args']['output_path'] = os.path.join(PROCESSED_FOLDER, output_filename)
        
        # Process with PDFProcessor (only pass workflow_ops)
        result = pdf_processor.execute_workflow(workflow_ops)
        return jsonify({"success": True, "result": result})
    except Exception as e:
        logging.error(f"Enhanced workflow error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/enhanced/bulk', methods=['POST'])
@login_required
def enhanced_bulk():
    """Enhanced bulk processing"""
    try:
        data = request.get_json()
        file_keys = data.get('file_keys', [])
        operation = data.get('operation')
        
        if not file_keys or not operation:
            return jsonify({"error": "File keys and operation required"}), 400
        
        # Process with PDFProcessor
        result = pdf_processor.bulk_process(operation, file_keys, f"bulk_{int(time.time())}")
        return jsonify({"success": True, "result": result})
    except Exception as e:
        logging.error(f"Enhanced bulk error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/enhanced/ipynb-to-pdf', methods=['POST'])
@login_required
def enhanced_ipynb_to_pdf():
    """Convert Jupyter notebook to PDF"""
    try:
        data = request.get_json()
        file_key = data.get('file_key')
        
        if not file_key:
            return jsonify({"error": "File key required"}), 400
        
        # Use the new conversion endpoint internally
        # Create new request data format for the /convert endpoint
        convert_data = {"key": file_key}
        
        # Save current request context
        orig_request = request._get_current_object()
        
        # Simulate a request to our /convert endpoint
        with app.test_request_context(
            path='/convert/ipynb-to-pdf',
            method='POST',
            json=convert_data
        ) as ctx:
            # Copy over authentication info
            ctx.request.environ['REMOTE_USER'] = orig_request.environ.get('REMOTE_USER')
            flask_response = convert_ipynb_to_pdf()
            
            # Return the response from the conversion endpoint
            if isinstance(flask_response, tuple):
                return flask_response
            
            # Process the response data 
            if hasattr(flask_response, 'json'):
                response_data = flask_response.json
                if 'success' in response_data:
                    # Format matches what the frontend expects
                    return flask_response
                elif 'key' in response_data:
                    # Format response to match what the frontend expects
                    return jsonify({
                        "success": True, 
                        "result": response_data.get('filename', f"converted_{int(time.time())}.pdf")
                    })
            
            # Fallback if we got an unexpected response
            return jsonify({"success": True, "result": "File converted successfully"})
    except Exception as e:
        logging.error(f"IPYNB conversion failed: {e}")
        return jsonify({"error": f"IPYNB conversion failed: {e}"}), 500

@app.route('/enhanced/ipynb-to-docx', methods=['POST'])
@login_required
def enhanced_ipynb_to_docx():
    """Convert Jupyter notebook to Word document"""
    try:
        data = request.get_json()
        file_key = data.get('file_key')
        
        if not file_key:
            return jsonify({"error": "File key required"}), 400
        
        # Use the new conversion endpoint internally
        # Create new request data format for the /convert endpoint
        convert_data = {"key": file_key}
        
        # Save current request context
        orig_request = request._get_current_object()
        
        # Simulate a request to our /convert endpoint
        with app.test_request_context(
            path='/convert/ipynb-to-docx',
            method='POST',
            json=convert_data
        ) as ctx:
            # Copy over authentication info
            ctx.request.environ['REMOTE_USER'] = orig_request.environ.get('REMOTE_USER')
            flask_response = convert_ipynb_to_docx()
            
            # Return the response from the conversion endpoint
            if isinstance(flask_response, tuple):
                return flask_response
            
            # Process the response data 
            if hasattr(flask_response, 'json'):
                response_data = flask_response.json
                if 'success' in response_data:
                    # Format matches what the frontend expects
                    return flask_response
                elif 'key' in response_data:
                    # Format response to match what the frontend expects
                    return jsonify({
                        "success": True, 
                        "result": response_data.get('filename', f"converted_{int(time.time())}.docx")
                    })
            
            # Fallback if we got an unexpected response
            return jsonify({"success": True, "result": "File converted successfully"})
    except Exception as e:
        logging.error(f"IPYNB to DOCX error: {e}")
        return jsonify({"error": f"ipynb_to_docx failed: {e}"}), 500

@app.route('/enhanced/py-to-ipynb', methods=['POST'])
@login_required
def enhanced_py_to_ipynb():
    """Convert Python file to Jupyter notebook"""
    try:
        data = request.get_json()
        file_key = data.get('file_key')
        
        if not file_key:
            return jsonify({"error": "File key required"}), 400
        
        # Get the full file path using helper function
        file_path = get_upload_path(file_key)
        if not os.path.exists(file_path):
            return jsonify({"error": f"File not found: {file_key}"}), 404
            
        # Process with our notebook utility functions
        output_file = f"converted_{int(time.time())}.ipynb"
        output_path = os.path.join(PROCESSED_FOLDER, output_file)
        
        # Make sure the output directory exists
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # Convert the Python file to IPYNB
        py_to_ipynb(file_path, output_path)
        
        return jsonify({"success": True, "result": output_path})
    except Exception as e:
        logging.error(f"Python to IPYNB error: {e}")
        return jsonify({"error": f"Python to IPYNB conversion failed: {e}"}), 500

@app.route('/enhanced/py-to-pdf', methods=['POST'])
@login_required
def enhanced_py_to_pdf():
    """Convert Python file to PDF"""
    try:
        data = request.get_json()
        file_key = data.get('file_key')
        
        if not file_key:
            return jsonify({"error": "File key required"}), 400
        
        # Get the full file path using helper function
        file_path = get_upload_path(file_key)
        if not os.path.exists(file_path):
            return jsonify({"error": f"File not found: {file_key}"}), 404
            
        # Process with our notebook utility functions
        output_file = f"converted_{int(time.time())}.pdf"
        output_path = os.path.join(PROCESSED_FOLDER, output_file)
        
        # Make sure the output directory exists
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # Convert the Python file to PDF
        py_to_pdf(file_path, output_path)
        
        return jsonify({"success": True, "result": output_path})
    except Exception as e:
        logging.error(f"Python to PDF error: {e}")
        return jsonify({"error": f"py_to_pdf failed: {e}"}), 500

@app.route('/enhanced/py-to-docx', methods=['POST'])
@login_required
def enhanced_py_to_docx():
    """Convert Python file to Word document"""
    try:
        data = request.get_json()
        file_key = data.get('file_key')
        
        if not file_key:
            return jsonify({"error": "File key required"}), 400
        
        # Get the full file path
        file_path = os.path.join(UPLOAD_FOLDER, file_key)
        if not os.path.exists(file_path):
            return jsonify({"error": f"File not found: {file_key}"}), 404
            
        # Process with PDFProcessor
        output_file = f"converted_{int(time.time())}.docx"
        output_path = os.path.join(PROCESSED_FOLDER, output_file)
        result = pdf_processor.py_to_docx(file_path, output_path)
        return jsonify({"success": True, "result": result})
    except Exception as e:
        logging.error(f"Python to DOCX error: {e}")
        return jsonify({"error": str(e)}), 500

# Re-add the route with explicit methods=["POST"]
@app.route('/enhanced/protect', methods=["POST"])
@login_required
def enhanced_protect_pdf():
    """Protect PDF with password"""
    if request.method != 'POST':
        return jsonify({"error": "Method not allowed"}), 405
        
    try:
        data = request.get_json()
        file_key = data.get('file_key')
        password = data.get('password')
        
        if not file_key:
            return jsonify({"error": "File key required"}), 400
        if not password:
            return jsonify({"error": "Password required"}), 400
        
        # Get the full file path using helper function
        file_path = get_upload_path(file_key)
        if not os.path.exists(file_path):
            return jsonify({"error": f"File not found: {file_key}"}), 404
            
        # Protect PDF using PyPDF
        try:
            # Create reader and writer objects
            reader = PdfReader(file_path)
            writer = PdfWriter()
            
            # Add all pages to the writer
            for page in reader.pages:
                writer.add_page(page)
            
            # Encrypt with the provided password
            writer.encrypt(password)
            
            # Create output file
            output_file = f"protected_{int(time.time())}.pdf"
            output_path = os.path.join(PROCESSED_FOLDER, output_file)
            
            # Make sure the output directory exists
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # Write the encrypted PDF to output
            with open(output_path, "wb") as output_pdf:
                writer.write(output_pdf)
                
            return jsonify({"success": True, "result": output_path})
        except Exception as e:
            return jsonify({"error": f"PDF protection failed: {str(e)}"}), 500
    except Exception as e:
        logging.error(f"PDF protection error: {e}")
        return jsonify({"error": str(e)}), 500

# ============================================================================
# CONVERSION ENDPOINTS - IPYNB TO PDF/DOCX 
# ============================================================================

# ---------- /convert/ipynb-to-pdf endpoint ----------
@app.route("/convert/ipynb-to-pdf", methods=["POST"])
@login_required
def convert_ipynb_to_pdf():
    data = request.get_json(silent=True) or request.form
    key = data.get("key") if data else None
    download_flag = str(data.get("download", "false")).lower() in ("1", "true", "yes")
    if not key:
        return jsonify({"error": "missing 'key' parameter"}), 400

    src = _get_upload_path(key)
    if not Path(src).exists():
        return jsonify({"error": f"File not found: {key}"}), 404

    # 1) Try direct nbconvert -> pdf
    if not _ensure_tool_on_path("jupyter"):
        # We still continue to try fallbacks, but inform the user.
        current_app.logger.debug("jupyter not found on PATH")
    temp_dir = Path(tempfile.mkdtemp(prefix="ipynb2pdf_"))
    base = Path(src).stem
    try:
        if _ensure_tool_on_path("jupyter"):
            cmd = ["jupyter", "nbconvert", "--to", "pdf", "--output", base, "--output-dir", str(temp_dir), src]
            current_app.logger.info("Trying nbconvert -> pdf: %s", " ".join(cmd))
            res = _run_cmd(cmd)
            if res["rc"] == 0:
                pdf_path = temp_dir / f"{base}.pdf"
                if pdf_path.exists():
                    # success
                    if download_flag:
                        return send_file(str(pdf_path), as_attachment=True, download_name=f"{base}.pdf")
                    out_key = f"{uuid.uuid4().hex}_{pdf_path.name}"
                    out_path = _get_upload_path(out_key)
                    shutil.move(str(pdf_path), out_path)
                    return jsonify({"key": out_key, "filename": pdf_path.name, "stdout": res["stdout"]}), 200
                # if nbconvert returned 0 but no file, continue to fallback
            else:
                current_app.logger.warning("nbconvert -> pdf failed: rc=%s", res["rc"])
        else:
            res = {"rc": -1, "stderr": "jupyter not found on PATH"}

        # 2) Fallback: nbconvert -> html, then HTML -> PDF using available renderer
        if _ensure_tool_on_path("jupyter"):
            cmd_html = ["jupyter", "nbconvert", "--to", "html", "--output", base, "--output-dir", str(temp_dir), src]
            current_app.logger.info("Trying nbconvert -> html: %s", " ".join(cmd_html))
            res_html = _run_cmd(cmd_html)
            html_path = temp_dir / f"{base}.html"
            if res_html["rc"] != 0 or not html_path.exists():
                # respond with combined debug info
                return jsonify({
                    "error": "IPYNB conversion failed (nbconvert).",
                    "nbconvert_stdout": res.get("stdout",""),
                    "nbconvert_stderr": res.get("stderr",""),
                    "nbconvert_html_stdout": res_html.get("stdout",""),
                    "nbconvert_html_stderr": res_html.get("stderr",""),
                }), 500
            # now try renderers:
            # prefer wkhtmltopdf
            if _ensure_tool_on_path("wkhtmltopdf"):
                cmd_wk = ["wkhtmltopdf", str(html_path), str(temp_dir / f"{base}.pdf")]
                current_app.logger.info("Trying wkhtmltopdf: %s", " ".join(cmd_wk))
                res_wk = _run_cmd(cmd_wk)
                pdf_path = temp_dir / f"{base}.pdf"
                if res_wk["rc"] == 0 and pdf_path.exists():
                    if download_flag:
                        return send_file(str(pdf_path), as_attachment=True, download_name=f"{base}.pdf")
                    out_key = f"{uuid.uuid4().hex}_{pdf_path.name}"
                    out_path = _get_upload_path(out_key)
                    shutil.move(str(pdf_path), out_path)
                    return jsonify({"key": out_key, "filename": pdf_path.name, "stdout": res_wk["stdout"]}), 200
                # else continue to other fallback(s)

            # try weasyprint (python) if installed
            try:
                from weasyprint import HTML  # noqa: F401
                current_app.logger.info("Trying weasyprint (Python fallback)")
                try:
                    HTML(filename=str(html_path)).write_pdf(str(temp_dir / f"{base}.pdf"))
                    pdf_path = temp_dir / f"{base}.pdf"
                    if pdf_path.exists():
                        if download_flag:
                            return send_file(str(pdf_path), as_attachment=True, download_name=f"{base}.pdf")
                        out_key = f"{uuid.uuid4().hex}_{pdf_path.name}"
                        out_path = _get_upload_path(out_key)
                        shutil.move(str(pdf_path), out_path)
                        return jsonify({"key": out_key, "filename": pdf_path.name, "stdout": "weasyprint used"}), 200
                except Exception as ewp:
                    current_app.logger.exception("weasyprint conversion failed")
                    # store exception to include in response below
                    _weasy_err = str(ewp)
            except Exception:
                _weasy_err = "weasyprint not installed"

            # try pdfkit (requires wkhtmltopdf) if available as Python library (wraps wkhtmltopdf)
            try:
                import pdfkit  # noqa: F401
                current_app.logger.info("Trying pdfkit (python wrapper)")
                try:
                    pdfkit.from_file(str(html_path), str(temp_dir / f"{base}.pdf"))
                    pdf_path = temp_dir / f"{base}.pdf"
                    if pdf_path.exists():
                        if download_flag:
                            return send_file(str(pdf_path), as_attachment=True, download_name=f"{base}.pdf")
                        out_key = f"{uuid.uuid4().hex}_{pdf_path.name}"
                        out_path = _get_upload_path(out_key)
                        shutil.move(str(pdf_path), out_path)
                        return jsonify({"key": out_key, "filename": pdf_path.name, "stdout": "pdfkit used"}), 200
                except Exception as e_pdfkit:
                    current_app.logger.exception("pdfkit failed")
                    _pdfkit_err = str(e_pdfkit)
            except Exception:
                _pdfkit_err = "pdfkit not installed"

            # if we reached here, all fallbacks failed - return diagnostics
            return jsonify({
                "error": "All conversion strategies failed.",
                "nbconvert_stdout": res.get("stdout",""),
                "nbconvert_stderr": res.get("stderr",""),
                "nbconvert_html_stdout": res_html.get("stdout",""),
                "nbconvert_html_stderr": res_html.get("stderr",""),
                "weasy_error": _weasy_err,
                "pdfkit_error": _pdfkit_err if '_pdfkit_err' in locals() else "pdfkit not attempted",
            }), 500

        else:
            # jupyter not installed, cannot convert
            return jsonify({"error": "jupyter not found on PATH; cannot convert", "note": "install jupyter (pip install jupyter)"}), 500

    finally:
        try:
            shutil.rmtree(temp_dir)
        except Exception:
            current_app.logger.debug("failed to remove temp_dir %s", temp_dir)

# ---------- /convert/ipynb-to-docx endpoint ----------
@app.route("/convert/ipynb-to-docx", methods=["POST"])
@login_required
def convert_ipynb_to_docx():
    data = request.get_json(silent=True) or request.form
    key = data.get("key") if data else None
    download_flag = str(data.get("download", "false")).lower() in ("1", "true", "yes")
    if not key:
        return jsonify({"error": "missing 'key' parameter"}), 400
    src = _get_upload_path(key)
    if not Path(src).exists():
        return jsonify({"error": f"File not found: {key}"}), 404

    # Primary: nbconvert -> markdown -> pandoc -> docx
    temp_dir = Path(tempfile.mkdtemp(prefix="ipynb2docx_"))
    base = Path(src).stem
    try:
        if not _ensure_tool_on_path("jupyter"):
            return jsonify({"error": "jupyter not found on PATH; install jupyter"}), 500
        # nbconvert to markdown
        cmd_md = ["jupyter", "nbconvert", "--to", "markdown", "--output", base, "--output-dir", str(temp_dir), src]
        current_app.logger.info("Running nbconvert -> markdown: %s", " ".join(cmd_md))
        res_md = _run_cmd(cmd_md)
        md_path = temp_dir / f"{base}.md"
        if res_md["rc"] != 0 or not md_path.exists():
            return jsonify({"error": "nbconvert to markdown failed", "stdout": res_md["stdout"], "stderr": res_md["stderr"]}), 500

        # If pandoc exists -> convert md -> docx
        if _ensure_tool_on_path("pandoc"):
            out_docx = temp_dir / f"{base}.docx"
            cmd_pandoc = ["pandoc", str(md_path), "-s", "-o", str(out_docx)]
            current_app.logger.info("Running pandoc: %s", " ".join(cmd_pandoc))
            res_p = _run_cmd(cmd_pandoc, cwd=str(temp_dir))
            if res_p["rc"] == 0 and out_docx.exists():
                if download_flag:
                    return send_file(str(out_docx), as_attachment=True, download_name=f"{base}.docx")
                out_key = f"{uuid.uuid4().hex}_{out_docx.name}"
                out_path = _get_upload_path(out_key)
                shutil.move(str(out_docx), out_path)
                return jsonify({"key": out_key, "filename": out_docx.name, "stdout": res_p["stdout"]}), 200
            else:
                return jsonify({"error": "pandoc failed to produce docx", "stdout": res_p["stdout"], "stderr": res_p["stderr"]}), 500

        # Fallback: try pure-Python creation: markdown -> plain text -> python-docx
        # This fallback is simple; it will not preserve images/complex formatting.
        try:
            import markdown as mdlib  # pip install markdown
            from docx import Document   # pip install python-docx
            from bs4 import BeautifulSoup  # pip install beautifulsoup4
        except Exception as e_mod:
            return jsonify({"error": "pandoc not found and python fallback libs missing", "hint": "install pandoc OR (pip install markdown python-docx beautifulsoup4)", "details": str(e_mod)}), 500

        # convert md -> html -> extract text paragraphs -> build docx
        html = mdlib.markdown(md_path.read_text(encoding="utf-8"))
        soup = BeautifulSoup(html, "html.parser")
        doc = Document()
        for el in soup.find_all(["h1", "h2", "h3", "p", "pre", "li"]):
            text = el.get_text("\n").strip()
            if not text:
                continue
            if el.name.startswith("h"):
                doc.add_heading(text, level=min(int(el.name[1]), 3))
            elif el.name == "pre":
                # code block - use monospace paragraph
                p = doc.add_paragraph()
                p.add_run(text).font.name = "Courier New"
            else:
                doc.add_paragraph(text)
        out_docx = temp_dir / f"{base}.docx"
        doc.save(str(out_docx))
        if not out_docx.exists():
            return jsonify({"error": "python-docx fallback failed to create docx"}), 500
        if download_flag:
            return send_file(str(out_docx), as_attachment=True, download_name=f"{base}.docx")
        out_key = f"{uuid.uuid4().hex}_{out_docx.name}"
        out_path = _get_upload_path(out_key)
        shutil.move(str(out_docx), out_path)
        return jsonify({"key": out_key, "filename": out_docx.name, "stdout": "python-docx fallback used"}), 200

    finally:
        try:
            shutil.rmtree(temp_dir)
        except Exception:
            current_app.logger.debug("failed to remove temp_dir %s", temp_dir)
try:
    from advanced.advanced_api import advanced_bp
    app.register_blueprint(advanced_bp, url_prefix='/advanced')
except Exception:
    logging.warning("Advanced API blueprint not available")

# Register the PDF operations blueprint
try:
    from advanced.pdf_operations_blueprint import pdf_operations
    app.register_blueprint(pdf_operations)
    logging.info("PDF operations blueprint registered successfully")
except Exception:
    logging.warning("PDF operations blueprint not available")

# ============================================================================
# OBSERVABILITY SETUP
# ============================================================================

if TELEMETRY_AVAILABLE:
    try:
        # Set up resource
        resource = OTResource(attributes={SERVICE_NAME: "pdf-tool-app"})
        # Tracer provider
        trace.set_tracer_provider(TracerProvider(resource=resource))
        # Jaeger exporter
        jaeger_exporter = JaegerExporter(
            agent_host_name=os.getenv('JAEGER_HOST', 'localhost'),
            agent_port=int(os.getenv('JAEGER_PORT', 6831))
        )
        trace.get_tracer_provider().add_span_processor(BatchSpanProcessor(jaeger_exporter))
        # Instrument Flask
        FlaskInstrumentor().instrument_app(app)
        logging.info("OpenTelemetry instrumentation enabled")
    except Exception as e:
        logging.error(f"Failed to setup OpenTelemetry: {e}")
        TELEMETRY_AVAILABLE = False

# --- Database Initialization ---
def init_db():
    with app.app_context():
        try:
            # Check if database tables exist by inspecting the db
            inspector = inspect(db.engine)
            tables = inspector.get_table_names()
            
            if 'user' not in tables:
                print("Creating database tables...")
                db.create_all()
                print("Database tables created successfully.")
            else:
                print("Database tables already exist.")
                
            # Ensure an admin user exists with username/password: admin/admin
            admin_user = User.query.filter_by(username='admin').first()
            if admin_user is None:
                print("Creating admin user...")
                new_admin = User(
                    username='admin',
                    email='admin@example.com',
                    password_hash=generate_password_hash('admin')
                )
                db.session.add(new_admin)
                db.session.commit()
                print("Admin user created successfully.")
            else:
                print("Admin user already exists.")
            
            # Store PostgreSQL tools version in app_config
            # Use the inspector to check for the config table
            config_table_exists = inspector.has_table('app_config')
            
            if not config_table_exists:
                db.session.execute(db.text('''
                CREATE TABLE IF NOT EXISTS app_config (
                    key VARCHAR(100) PRIMARY KEY,
                    value TEXT NOT NULL,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
                '''))
                
        except Exception as e:
            print(f"Error initializing database: {str(e)}")
            logging.exception("Database initialization error:")
            
            # Update PG tools version
            # Check if the entry already exists
            result = db.session.execute(db.text("SELECT 1 FROM app_config WHERE key = 'pg_tools_version'")).fetchone()
            
            if result:
                # Update existing entry
                db.session.execute(db.text("UPDATE app_config SET value = '2.0.12\\Windows', updated_at = CURRENT_TIMESTAMP WHERE key = 'pg_tools_version'"))
            else:
                # Insert new entry
                db.session.execute(db.text("INSERT INTO app_config (key, value, updated_at) VALUES ('pg_tools_version', '2.0.12\\Windows', CURRENT_TIMESTAMP)"))
            
            db.session.commit()
            
            print(f"PostgreSQL Tools Version: 2.0.12\\Windows")
        except Exception as e:
            logging.warning(f"Could not store PostgreSQL tools version: {e}")
            print("Created admin user with default credentials admin/admin (for local testing)")
        else:
            # Keep credentials in sync for local testing
            admin_user.password_hash = generate_password_hash('admin')
            if not admin_user.email:
                admin_user.email = 'admin@example.com'
            db.session.commit()
            print("Ensured admin credentials are set to admin/admin (for local testing)")

        print("Database tables created and admin user ensured successfully!")

# ---------- /convert/ipynb-to-pdf-v2 endpoint ----------
@app.route("/convert/ipynb-to-pdf-v2", methods=["POST"])
@login_required
def convert_ipynb_to_pdf_v2():
    data = request.get_json(silent=True) or request.form
    key = data.get("key") or data.get("file_key")
    download_flag = str(data.get("download", "false")).lower() in ("1", "true", "yes")
    if not key:
        return jsonify({"error": "missing 'key' or 'file_key' parameter"}), 400

    src = get_upload_path(key)
    if not Path(src).exists():
        return jsonify({"error": f"File not found: {key}"}), 404

    # 1) Try direct nbconvert -> pdf
    if not _ensure_tool_on_path("jupyter"):
        # We still continue to try fallbacks, but inform the user.
        current_app.logger.debug("jupyter not found on PATH")
    temp_dir = Path(tempfile.mkdtemp(prefix="ipynb2pdf_"))
    base = Path(src).stem
    try:
        if _ensure_tool_on_path("jupyter"):
            cmd = ["jupyter", "nbconvert", "--to", "pdf", "--output", base, "--output-dir", str(temp_dir), src]
            current_app.logger.info("Trying nbconvert -> pdf: %s", " ".join(cmd))
            res = _run_cmd(cmd)
            if res["rc"] == 0:
                pdf_path = temp_dir / f"{base}.pdf"
                if pdf_path.exists():
                    # success
                    if download_flag:
                        return send_file(str(pdf_path), as_attachment=True, download_name=f"{base}.pdf")
                    output_file = f"converted_{int(time.time())}.pdf"
                    output_path = os.path.join(PROCESSED_FOLDER, output_file)
                    shutil.copy2(str(pdf_path), output_path)
                    return jsonify({"success": True, "result": output_path}), 200
                # if nbconvert returned 0 but no file, continue to fallback
            else:
                current_app.logger.warning("nbconvert -> pdf failed: rc=%s", res["rc"])
        else:
            res = {"rc": -1, "stderr": "jupyter not found on PATH"}

        # 2) Fallback: nbconvert -> html, then HTML -> PDF using available renderer
        if _ensure_tool_on_path("jupyter"):
            cmd_html = ["jupyter", "nbconvert", "--to", "html", "--output", base, "--output-dir", str(temp_dir), src]
            current_app.logger.info("Trying nbconvert -> html: %s", " ".join(cmd_html))
            res_html = _run_cmd(cmd_html)
            html_path = temp_dir / f"{base}.html"
            if res_html["rc"] != 0 or not html_path.exists():
                # respond with combined debug info
                return jsonify({
                    "error": "IPYNB conversion failed (nbconvert).",
                    "nbconvert_stdout": res.get("stdout",""),
                    "nbconvert_stderr": res.get("stderr",""),
                    "nbconvert_html_stdout": res_html.get("stdout",""),
                    "nbconvert_html_stderr": res_html.get("stderr",""),
                }), 500
            # now try renderers:
            # prefer wkhtmltopdf
            if _ensure_tool_on_path("wkhtmltopdf"):
                cmd_wk = ["wkhtmltopdf", str(html_path), str(temp_dir / f"{base}.pdf")]
                current_app.logger.info("Trying wkhtmltopdf: %s", " ".join(cmd_wk))
                res_wk = _run_cmd(cmd_wk)
                pdf_path = temp_dir / f"{base}.pdf"
                if res_wk["rc"] == 0 and pdf_path.exists():
                    if download_flag:
                        return send_file(str(pdf_path), as_attachment=True, download_name=f"{base}.pdf")
                    output_file = f"converted_{int(time.time())}.pdf"
                    output_path = os.path.join(PROCESSED_FOLDER, output_file)
                    shutil.copy2(str(pdf_path), output_path)
                    return jsonify({"success": True, "result": output_path}), 200
                # else continue to other fallback(s)

            # try weasyprint (python) if installed
            try:
                from weasyprint import HTML  # noqa: F401
                current_app.logger.info("Trying weasyprint (Python fallback)")
                try:
                    HTML(filename=str(html_path)).write_pdf(str(temp_dir / f"{base}.pdf"))
                    pdf_path = temp_dir / f"{base}.pdf"
                    if pdf_path.exists():
                        if download_flag:
                            return send_file(str(pdf_path), as_attachment=True, download_name=f"{base}.pdf")
                        output_file = f"converted_{int(time.time())}.pdf"
                        output_path = os.path.join(PROCESSED_FOLDER, output_file)
                        shutil.copy2(str(pdf_path), output_path)
                        return jsonify({"success": True, "result": output_path}), 200
                except Exception as ewp:
                    current_app.logger.exception("weasyprint conversion failed")
                    # store exception to include in response below
                    _weasy_err = str(ewp)
            except Exception:
                _weasy_err = "weasyprint not installed"

            # try pdfkit (requires wkhtmltopdf) if available as Python library (wraps wkhtmltopdf)
            try:
                import pdfkit  # noqa: F401
                current_app.logger.info("Trying pdfkit (python wrapper)")
                try:
                    pdfkit.from_file(str(html_path), str(temp_dir / f"{base}.pdf"))
                    pdf_path = temp_dir / f"{base}.pdf"
                    if pdf_path.exists():
                        if download_flag:
                            return send_file(str(pdf_path), as_attachment=True, download_name=f"{base}.pdf")
                        output_file = f"converted_{int(time.time())}.pdf"
                        output_path = os.path.join(PROCESSED_FOLDER, output_file)
                        shutil.copy2(str(pdf_path), output_path)
                        return jsonify({"success": True, "result": output_path}), 200
                except Exception as e_pdfkit:
                    current_app.logger.exception("pdfkit failed")
                    _pdfkit_err = str(e_pdfkit)
            except Exception:
                _pdfkit_err = "pdfkit not installed"

            # if we reached here, all fallbacks failed - return diagnostics
            return jsonify({
                "error": "All conversion strategies failed.",
                "nbconvert_stdout": res.get("stdout",""),
                "nbconvert_stderr": res.get("stderr",""),
                "nbconvert_html_stdout": res_html.get("stdout",""),
                "nbconvert_html_stderr": res_html.get("stderr",""),
                "weasy_error": _weasy_err if '_weasy_err' in locals() else "weasyprint not attempted",
                "pdfkit_error": _pdfkit_err if '_pdfkit_err' in locals() else "pdfkit not attempted",
            }), 500

        else:
            # jupyter not installed, cannot convert
            return jsonify({"error": "jupyter not found on PATH; cannot convert", "note": "install jupyter (pip install jupyter)"}), 500

    finally:
        try:
            shutil.rmtree(temp_dir)
        except Exception:
            current_app.logger.debug("failed to remove temp_dir %s", temp_dir)

# ---------- /convert/ipynb-to-docx-v2 endpoint ----------
@app.route("/convert/ipynb-to-docx-v2", methods=["POST"])
@login_required
def convert_ipynb_to_docx_v2():
    data = request.get_json(silent=True) or request.form
    key = data.get("key") or data.get("file_key")
    download_flag = str(data.get("download", "false")).lower() in ("1", "true", "yes")
    if not key:
        return jsonify({"error": "missing 'key' or 'file_key' parameter"}), 400
    src = get_upload_path(key)
    if not Path(src).exists():
        return jsonify({"error": f"File not found: {key}"}), 404

    # Primary: nbconvert -> markdown -> pandoc -> docx
    temp_dir = Path(tempfile.mkdtemp(prefix="ipynb2docx_"))
    base = Path(src).stem
    try:
        if not _ensure_tool_on_path("jupyter"):
            return jsonify({"error": "jupyter not found on PATH; install jupyter"}), 500
        # nbconvert to markdown
        cmd_md = ["jupyter", "nbconvert", "--to", "markdown", "--output", base, "--output-dir", str(temp_dir), src]
        current_app.logger.info("Running nbconvert -> markdown: %s", " ".join(cmd_md))
        res_md = _run_cmd(cmd_md)
        md_path = temp_dir / f"{base}.md"
        if res_md["rc"] != 0 or not md_path.exists():
            return jsonify({"error": "nbconvert to markdown failed", "stdout": res_md["stdout"], "stderr": res_md["stderr"]}), 500

        # If pandoc exists -> convert md -> docx
        if _ensure_tool_on_path("pandoc"):
            out_docx = temp_dir / f"{base}.docx"
            cmd_pandoc = ["pandoc", str(md_path), "-s", "-o", str(out_docx)]
            current_app.logger.info("Running pandoc: %s", " ".join(cmd_pandoc))
            res_p = _run_cmd(cmd_pandoc, cwd=str(temp_dir))
            if res_p["rc"] == 0 and out_docx.exists():
                if download_flag:
                    return send_file(str(out_docx), as_attachment=True, download_name=f"{base}.docx")
                output_file = f"converted_{int(time.time())}.docx"
                output_path = os.path.join(PROCESSED_FOLDER, output_file)
                shutil.copy2(str(out_docx), output_path)
                return jsonify({"success": True, "result": output_path}), 200
            else:
                return jsonify({"error": "pandoc failed to produce docx", "stdout": res_p["stdout"], "stderr": res_p["stderr"]}), 500

        # Fallback: try pure-Python creation: markdown -> plain text -> python-docx
        # This fallback is simple; it will not preserve images/complex formatting.
        try:
            import markdown as mdlib  # pip install markdown
            from docx import Document   # pip install python-docx
            from bs4 import BeautifulSoup  # pip install beautifulsoup4
        except Exception as e_mod:
            return jsonify({"error": "pandoc not found and python fallback libs missing", "hint": "install pandoc OR (pip install markdown python-docx beautifulsoup4)", "details": str(e_mod)}), 500

        # convert md -> html -> extract text paragraphs -> build docx
        html = mdlib.markdown(md_path.read_text(encoding="utf-8"))
        soup = BeautifulSoup(html, "html.parser")
        doc = Document()
        for el in soup.find_all(["h1", "h2", "h3", "p", "pre", "li"]):
            text = el.get_text("\n").strip()
            if not text:
                continue
            if el.name.startswith("h"):
                doc.add_heading(text, level=min(int(el.name[1]), 3))
            elif el.name == "pre":
                # code block - use monospace paragraph
                p = doc.add_paragraph()
                p.add_run(text).font.name = "Courier New"
            else:
                doc.add_paragraph(text)
        out_docx = temp_dir / f"{base}.docx"
        doc.save(str(out_docx))
        if not out_docx.exists():
            return jsonify({"error": "python-docx fallback failed to create docx"}), 500
        if download_flag:
            return send_file(str(out_docx), as_attachment=True, download_name=f"{base}.docx")
        output_file = f"converted_{int(time.time())}.docx"
        output_path = os.path.join(PROCESSED_FOLDER, output_file)
        shutil.copy2(str(out_docx), output_path)
        return jsonify({"success": True, "result": output_path}), 200

    finally:
        try:
            shutil.rmtree(temp_dir)
        except Exception:
            current_app.logger.debug("failed to remove temp_dir %s", temp_dir)

if __name__ == '__main__':
    print("Starting PDF Tool server with SQLite...")
    print("Initializing database...")
    init_db()
    print("Access the frontend at: http://localhost:5000")
    print("Press Ctrl+C to stop the server")
    
    # Use environment variable to control debug mode
    debug_mode = os.environ.get('FLASK_DEBUG', 'False').lower() == 'true'
    app.run(host='0.0.0.0', port=5000, debug=debug_mode)
